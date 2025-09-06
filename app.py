APP_TITLE = "MRIDAASTRO"
APP_TAGLINE = "In the light of divine, let your soul journey shine"

# ===== Background Template Helper (stable image) =====
import os
from io import BytesIO
from docx import Document as _WordDocument
from docx.shared import RGBColor

TEMPLATE_DOCX = "bg_template.docx"

def make_document():
    try:
        if os.path.exists(TEMPLATE_DOCX):
            return _WordDocument(TEMPLATE_DOCX)
    except Exception:
        pass
    return _WordDocument()
# ===== End Background Template Helper =====


# app_docx_borders_85pt_editable_v6_8_8_locked.py
# Changes from 6.8.7:
# - Rename & style headings:
#     * "Planetary Positions..." -> "ग्रह स्थिति" (bold + underline)
#     * "Vimshottari Mahadasha..." -> "विंशोत्तरी महादशा" (bold + underline)
# - Fix kundali preview image whitespace: compact square PNG with zero padding

import math
import datetime, json, urllib.parse, urllib.request
from io import BytesIO

# --- One-page layout switch ---
ONE_PAGE = True

# --- Appearance configuration ---
# Sizing (pt) — tuned smaller to reduce white space
NUM_W_PT = 10       # house number box width (was 12)
NUM_H_PT = 12       # house number box height (was 14)
PLANET_W_PT = 20    # planet label box width (was 16)
PLANET_H_PT = 16    # planet label box height (was 14)
GAP_X_PT = 3        # horizontal gap between planet boxes (was 4)
OFFSET_Y_PT = 10    # vertical offset below number box (was 12)

# ===== MODERN CHART STYLING OPTIONS =====
# Options: "plain", "bordered", "shaded", "bordered_shaded"
HOUSE_NUM_STYLE = "bordered_shaded"
HOUSE_NUM_BORDER_PT = 1.0
HOUSE_NUM_SHADE = "#f8f9fa"  # Light gray for modern look

# Modern color scheme for charts
CHART_COLORS = {
    'house_border': '#194A6D',     # Deep blue
    'house_fill': '#f8f9fa',      # Light gray
    'planet_benefic': '#2E8B57',  # Sea green for benefic planets
    'planet_malefic': '#DC143C',  # Crimson for malefic planets
    'planet_neutral': '#4682B4',  # Steel blue for neutral planets
    'number_bg': '#ffffff',       # White for house numbers
    'text_primary': '#2d3748',    # Dark gray for text
}




# --- Reliable cell shading (works in all Word views) ---
def shade_cell(cell, fill_hex="FFFFFF"):
    return

def shade_header_row(table, fill_hex="FFFFFF"):
    return


def compact_document_spacing(doc):
    """Reduce vertical whitespace across the document."""
    try:
        from docx.shared import Pt
        try:
            st = doc.styles["Normal"].paragraph_format
            st.space_before = Pt(0)
            st.space_after = Pt(0)
            st.line_spacing = 1.0
        except Exception:
            pass
        for p in doc.paragraphs:
            try:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        try:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                        except Exception:
                            pass
    except Exception:
        pass
def set_page_background(doc, hex_color):
    try:
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), hex_color)
        doc.element.insert(0, bg)
    except Exception:
        pass


# --- Phalit ruled lines (25 rows) ---
from docx.enum.table import WD_ROW_HEIGHT_RULE

def zero_table_cell_margins(table):
    """Set w:tblCellMar for all sides to 0 to remove extra top/bottom padding inside table cells."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        # Remove existing cell margins if present
        for el in list(tblPr):
            if el.tag.endswith('tblCellMar'):
                tblPr.remove(el)
        cellMar = OxmlElement('w:tblCellMar')
        for side in ('top','left','bottom','right'):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '0')
            m.set(qn('w:type'), 'dxa')
            cellMar.append(m)
        tblPr.append(cellMar)
    except Exception:
        pass
def add_phalit_section(container_cell, width_inches=3.60, rows=15):
    # Add beautiful cylindrical gradient header bar for फलित section
    create_cylindrical_section_header(container_cell, "फलित", width_pt=260)

    t = container_cell.add_table(rows=rows, cols=1); t.autofit = False
    # Clear table borders so only bottom rules show
    try:
        tbl = t._tbl; tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'nil'); tblBorders.append(el)
        tblPr.append(tblBorders)
    except Exception:
        pass
    set_col_widths(t, [width_inches])
    for r in t.rows:
        r.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        r.height = Pt(14)
        c = r.cells[0]
        p = c.paragraphs[0]; run = p.add_run("\u00A0"); run.font.size = Pt(1)
        tcPr = c._tc.get_or_add_tcPr()
        for el in list(tcPr):
            if el.tag.endswith('tcBorders'):
                tcPr.remove(el)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top','left','right'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'nil'); tcBorders.append(el)
        el = OxmlElement('w:bottom')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8'); el.set(qn('w:space'),'0'); el.set(qn('w:color'),'E67E22')
        tcBorders.append(el)
        tcPr.append(tcBorders)

def _rects_overlap(a, b):
    return not (a['right'] <= b['left'] or a['left'] >= b['right'] or a['bottom'] <= b['top'] or a['top'] >= b['bottom'])

def _nudge_number_box(base_left, base_top, w, h, S, occupied):
    cx = S/2.0; cy = S/2.0
    bx = base_left + w/2.0; by = base_top + h/2.0
    vx = (bx - cx); vy = (by - cy)
    n = (vx*vx + vy*vy) ** 0.5 or 1.0
    ux, uy = vx/n, vy/n  # unit vector outward
    pad = 2.0
    for step in range(0, 9):  # try nudges up to ~16pt
        dx = ux * (step * 2.0)
        dy = uy * (step * 2.0)
        l = max(pad, min(S - w - pad, base_left + dx))
        t = max(pad, min(S - h - pad, base_top + dy))
        r = {'left': l, 'top': t, 'right': l + w, 'bottom': t + h}
        hit = False
        for o in occupied:
            if _rects_overlap(r, o):
                hit = True; break
        if not hit:
            return l, t
    return base_left, base_top
import matplotlib.pyplot as plt
import pandas as pd
import pytz
import streamlit as st
# === App background helper (for authenticated pages) ===
import base64, os, streamlit as st

def set_app_background(image_path: str, size: str = "contain", position: str = "top center"):
    """
    Shows a page background image on *authenticated* pages.
    Call this after login/whitelist check, before rendering the UI.
    """
    try:
        if not os.path.exists(image_path):
            return
        with open(image_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        st.markdown(f"""
        <style>
          [data-testid="stAppViewContainer"] {{
            background-image: url("data:image/png;base64,{b64}");
            background-size: {size};
            background-position: {position};
            background-repeat: no-repeat;
            background-color: #f6ede6; /* fallback */
          }}
          [data-testid="stHeader"] {{
            background: transparent;
          }}
        </style>
        """, unsafe_allow_html=True)
    except Exception:
        pass
# === End background helper ===


# --- Google Search Console verification (inject into <head>) ---
st.markdown("""
<script>
(function() {
  try {
    var meta = document.createElement('meta');
    meta.name = 'google-site-verification';
    meta.content = '01pSw-vPDjcZLjPluDXzbWvMR-YxFjh3w3T94nMxsVU';
    document.getElementsByTagName('head')[0].appendChild(meta);
  } catch (e) { console.log('GSC meta inject error', e); }
})();
</script>
""", unsafe_allow_html=True)


from login_branding_helper import show_login_screen

# ===================== Google OAuth2 Login Gate (with callback) =====================
import time, requests
from urllib.parse import urlencode
from google.oauth2 import id_token
from google.auth.transport import requests as g_requests
import streamlit as st

# Read secrets (supports both top-level and [google_oauth] section)
try:
    _cfg = st.secrets.get("google_oauth", st.secrets)
    CLIENT_ID     = _cfg["client_id"]
    CLIENT_SECRET = _cfg["client_secret"]
    REDIRECT_URI  = _cfg["redirect_uri"]  # e.g. https://mridaastro.streamlit.app/oauth2callback
    OAUTH_ENABLED = True
except:
    # Demo mode - OAuth not configured
    CLIENT_ID     = "demo"
    CLIENT_SECRET = "demo"
    REDIRECT_URI  = "demo"
    OAUTH_ENABLED = False

AUTH_ENDPOINT  = "https://accounts.google.com/o/oauth2/v2/auth"
TOKEN_ENDPOINT = "https://oauth2.googleapis.com/token"
SCOPES = "openid email profile"

def build_auth_url(state: str) -> str:
    params = {
        "client_id": CLIENT_ID,
        "redirect_uri": REDIRECT_URI,
        "response_type": "code",
        "scope": SCOPES,
        "access_type": "online",
        "include_granted_scopes": "true",
        "prompt": "consent",
        "state": state,
    }
    return f"{AUTH_ENDPOINT}?{urlencode(params)}"

def exchange_code_for_tokens(code: str) -> dict:
    data = {
        "code": code,
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "redirect_uri": REDIRECT_URI,
        "grant_type": "authorization_code",
    }
    resp = requests.post(TOKEN_ENDPOINT, data=data, timeout=15)
    resp.raise_for_status()
    return resp.json()

def verify_id_token(idt: str) -> dict:
    # Verifies signature & audience (CLIENT_ID)
    return id_token.verify_oauth2_token(idt, g_requests.Request(), CLIENT_ID)

def sign_out():
    for k in ("user", "oauth", "oauth_state"):
        st.session_state.pop(k, None)
    st.rerun()

# --- Handle Google redirect (works on /oauth2callback or any path with ?code=...)
#     Use new stable API: st.query_params (replaces deprecated experimental_get_query_params)
qp = dict(st.query_params)  # convert to a plain dict
code  = qp.get("code")
state = qp.get("state")

# If values are lists, take the first element
if isinstance(code, list):
    code = code[0] if code else None
if isinstance(state, list):
    state = state[0] if state else None

if code:
    try:
        if "oauth_state" in st.session_state and state != st.session_state["oauth_state"]:
            st.error("State mismatch. Please try signing in again.")
            st.stop()

        tokens = exchange_code_for_tokens(code)
        claims = verify_id_token(tokens["id_token"])
        st.session_state["user"] = {
            "email": claims.get("email"),
            "name": claims.get("name") or claims.get("email"),
            "picture": claims.get("picture", ""),
        }
        st.session_state["oauth"] = tokens

        # Clear query params and send user back to root path
        st.query_params.clear()
        st.markdown("<script>history.replaceState({}, '', '/');</script>", unsafe_allow_html=True)

        st.success(f"Signed in as {st.session_state['user']['email']}")
        time.sleep(0.5)
        st.rerun()
    except Exception:
        st.error("Login failed. Please try again.")
        st.stop()

# --- If not signed in, show login and stop
if "user" not in st.session_state:
    # Render the fully branded login screen (background + hero + gold button)
    show_login_screen()
    st.stop()

# --- Restrict who can access (pick ONE approach) ---

# --- Restrict who can access (STRICT WHITELIST) ---
email = (st.session_state["user"].get("email") or "").lower()

# Read allowed users from Streamlit secrets. Supports either a list or a comma-separated string.
_allowed_raw = st.secrets.get("allowed_users", [])
if isinstance(_allowed_raw, str):
    allowed_users = {u.strip().lower() for u in _allowed_raw.split(",") if u.strip()}
elif isinstance(_allowed_raw, (list, tuple, set)):
    allowed_users = {str(u).strip().lower() for u in _allowed_raw if str(u).strip()}
else:
    allowed_users = set()

# Enforce: if whitelist is empty -> deny by default to avoid accidental exposure.
if not allowed_users:
    st.error("Access restricted. No allowed users configured. Add 'allowed_users' in Streamlit Secrets.")
    st.stop()

if email not in allowed_users:
    st.error("Access restricted to authorized users only.")
    st.stop()

# Set background for authenticated app pages
set_app_background("assets/login_bg.png", size="contain", position="top center")


# Show identity & Sign out in sidebar
st.sidebar.markdown(f"**Signed in:** {st.session_state['user'].get('name') or email} ({email})")
if st.sidebar.button("Sign out"):
    sign_out()
# =================== End Google OAuth2 Login Gate (with callback) ===================

# --- Custom style for Generate & Download buttons ---
st.markdown("""
    <style>
    div.stButton > button,
    div.stDownloadButton > button {
        background-color: black;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        border: 1px solid #2e8b57;
    }
    div.stButton > button:hover,
    div.stDownloadButton > button:hover {
        background-color: #2e8b57 !important;  /* sea green hover */
        color: white !important;
    }
    div.stButton > button:active,
    div.stDownloadButton > button:active {
        background-color: #2e8b57 !important;  /* sea green click */
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

from PIL import Image


# === App background (minimal, no logic changes) ===
def _apply_bg():
    try:
        import streamlit as st, base64
        from pathlib import Path
        p = Path("assets/ganesha_bg.png")
        if p.exists():
            b64 = base64.b64encode(p.read_bytes()).decode()
            css = f"""
            <style>
            [data-testid="stAppViewContainer"] {{
                background: url('data:image/png;base64,{b64}') no-repeat center top fixed;
                background-size: cover;
            }}
            </style>
            """
            st.markdown(css, unsafe_allow_html=True)
    except Exception:
        pass
# === End App background ===


import swisseph as swe
from timezonefinder import TimezoneFinder


def _bbox_of_poly(poly):
    xs, ys = zip(*poly)
    return {'left': min(xs), 'top': min(ys), 'right': max(xs), 'bottom': max(ys)}

def _clamp_in_bbox(left, top, w, h, bbox, pad):
    lmin = bbox['left'] + pad
    tmin = bbox['top'] + pad
    lmax = bbox['right'] - w - pad
    tmax = bbox['bottom'] - h - pad
    return max(lmin, min(left, lmax)), max(tmin, min(top, tmax))
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

# --- Table header shading helper (match kundali bg) ---
def shade_cell(cell, fill_hex="FFFFFF"):
    return

def shade_header_row(table, fill_hex="FFFFFF"):
    return

def set_page_background(doc, hex_color):
    try:
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), hex_color)
        doc.element.insert(0, bg)
    except Exception:
        pass



# ---- Dasha helpers (top-level; ORDER & YEARS must exist at call time) ----
def antar_segments_in_md_utc(md_lord, md_start_utc, md_days):
    res=[]; t=md_start_utc; start_idx=ORDER.index(md_lord)
    for i in range(9):
        L=ORDER[(start_idx+i)%9]; dur = YEARS[L]*(md_days/(120.0)); start = t; end = t + datetime.timedelta(days=dur)
        res.append((L, start, end, dur)); t = end
    return res

def pratyantars_in_antar_utc(antar_lord, antar_start_utc, antar_days):
    res=[]; t=antar_start_utc; start_idx=ORDER.index(antar_lord)
    for i in range(9):
        L=ORDER[(start_idx+i)%9]; dur = YEARS[L]*(antar_days/(120.0)); start = t; end = t + datetime.timedelta(days=dur)
        res.append((L, start, end)); t = end
    return res

def next_antar_in_days_utc(now_utc, md_segments, days_window):
    rows=[]; horizon=now_utc + datetime.timedelta(days=days_window)
    for seg in md_segments:
        MD = seg["planet"]; ms = seg["start"]; me = seg["end"]; md_days = seg["days"]
        for AL, as_, ae, adays in antar_segments_in_md_utc(MD, ms, md_days):
            if ae < now_utc or as_ > horizon: 
                continue
            end = min(ae, horizon)
            rows.append({"major": MD, "antar": AL, "end": end})
    rows.sort(key=lambda r:r["end"])
    return rows
# ---- End helpers ----


# ---- Dasha helpers (top-level; use ORDER & YEARS defined before calls) ----
# ---- End helpers ----



# --- favicon helper (must be defined before set_page_config) ---
def _load_page_icon():
    try:
        return Image.open("assets/fevicon_icon.png")
    except Exception:
        return "🪔"
st.set_page_config(page_title="MRIDAASTRO", layout="wide", page_icon=_load_page_icon())

# --- First-visit reset so 'Required' doesn't show on initial load ---
if 'first_visit' not in st.session_state:
    st.session_state['first_visit'] = True
if st.session_state.get('first_visit', False):
    st.session_state['submitted'] = False
    st.session_state['first_visit'] = False






# --- show validation only after first submit ---
if 'submitted' not in st.session_state:
    st.session_state['submitted'] = False

def render_label(text: str, show_required: bool = False):
    if show_required:
        # Professional validation: black field name + red "Required"
        html = (
            "<div style='display:flex;justify-content:space-between;align-items:center;'>"
            f"<span style='font-weight:700; font-size:18px;'>{text}</span>"
            "<span style='color:#c1121f; font-size:14px; font-weight:700;'>Required</span>"
            "</div>"
        )
    else:
        # Normal state
        html = (
            "<div style='display:flex;justify-content:space-between;align-items:center;'>"
            f"<span style='font-weight:700; font-size:18px;'>{text}</span>"
            "</div>"
        )
    st.markdown(html, unsafe_allow_html=True)

# === MRIDAASTRO Brand Header (Top) ===
st.markdown(
    """
    <link href="https://fonts.googleapis.com/css2?family=Cinzel+Decorative:wght@700&display=swap" rel="stylesheet">
    <div style='text-align:center; padding: 14px 0 4px 0;'>
      <div style='font-family:Cinzel Decorative, cursive; font-size:58px; font-weight:700; color:#000000; text-shadow:2px 2px 4px rgba(0,0,0,0.2); margin-bottom:8px;'>
        MRIDAASTRO
      </div>
      <div style='font-family:Georgia,serif; font-style:italic; font-size:24px; color:#000000; margin-bottom:18px;'>
        In the light of divine, let your soul journey shine
      </div>
      <div style='height:3px; width:160px; margin:0 auto 6px auto; background:black; border-radius:2px;'></div>
    </div>
    """,
    unsafe_allow_html=True
)
# === End MRIDAASTRO Header ===
_apply_bg()
AYANAMSHA_VAL = swe.SIDM_LAHIRI
YEAR_DAYS     = 365.2422

BASE_FONT_PT = 7.0
LATIN_FONT = "Georgia"
HINDI_FONT = "Mangal"

HN = {'Su':'सूर्य','Mo':'चंद्र','Ma':'मंगल','Me':'बुध','Ju':'गुरु','Ve':'शुक्र','Sa':'शनि','Ra':'राहु','Ke':'केतु'}

# Compact Hindi abbreviations for planet boxes
HN_ABBR = {'Su':'सू','Mo':'चं','Ma':'मं','Me':'बु','Ju':'गु','Ve':'शु','Sa':'श','Ra':'रा','Ke':'के'}

# ==== Status helpers (Rāśi vs Navāṁśa aware) ====
SIGN_LORD = {1:'Ma',2:'Ve',3:'Me',4:'Mo',5:'Su',6:'Me',7:'Ve',8:'Ma',9:'Ju',10:'Sa',11:'Sa',12:'Ju'}
EXALT_SIGN = {'Su':1,'Mo':2,'Ma':10,'Me':6,'Ju':4,'Ve':12,'Sa':7,'Ra':2,'Ke':8}
DEBIL_SIGN = {'Su':7,'Mo':8,'Ma':4,'Me':12,'Ju':10,'Ve':6,'Sa':1,'Ra':8,'Ke':2}
# --- Combustion settings ---
# Only the SUN causes combustion. Rahu/Ketu never combust. Moon CAN be combust (by Sun) if within orb.
# Set this to True if you want to mark combustion ONLY when the Sun and the planet are in the SAME rāśi sign.
REQUIRE_SAME_SIGN_FOR_COMBUST = False  # change to True if that matches your tradition

COMBUST_ORB = {'Mo':12.0,'Ma':17.0,'Me':12.0,'Ju':11.0,'Ve':10.0,'Sa':15.0}

def _min_circ_angle(a, b):
    d = abs((a - b) % 360.0)
    return d if d <= 180.0 else 360.0 - d

def _xml_text(s):
    return (str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))

def planet_rasi_sign(lon_sid):
    return int(lon_sid // 30) + 1  # 1..12

def compute_statuses_all(sidelons):
    """Return per-planet dict containing both rasi-based and nav-based flags."""
    out = {}
    sun_lon = sidelons.get('Su', 0.0)
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        lon = sidelons[code]
        rasi = planet_rasi_sign(lon)
        nav  = navamsa_sign_from_lon_sid(lon)
        varg = (rasi == nav)
        # Combustion: Sun only, optional same-sign constraint
        combust = False
        if code in COMBUST_ORB and code != 'Su':
            sep = _min_circ_angle(lon, sun_lon)
            if not REQUIRE_SAME_SIGN_FOR_COMBUST or (planet_rasi_sign(lon) == planet_rasi_sign(sun_lon)):
                combust = (sep <= COMBUST_ORB[code])

        out[code] = {
            'rasi': rasi,
            'nav': nav,
            'vargottama': varg,
            'combust': combust,
            'self_rasi': (SIGN_LORD.get(rasi) == code),
            'self_nav':  (SIGN_LORD.get(nav)  == code),
            'exalt_rasi': (EXALT_SIGN.get(code) == rasi),
            'exalt_nav':  (EXALT_SIGN.get(code) == nav),
            'debil_rasi': (DEBIL_SIGN.get(code) == rasi),
            'debil_nav':  (DEBIL_SIGN.get(code) == nav),
        }
        # Nodes (Rahu/Ketu): do not mark exaltation/debilitation
        if code in ('Ra','Ke'):
            out[code]['exalt_rasi'] = False
            out[code]['exalt_nav'] = False
            out[code]['debil_rasi'] = False
            out[code]['debil_nav'] = False
    return out

def _make_flags(view, st):
    """Reduce the big dict to the fields used by the renderer for a given chart view."""
    if view == 'nav':
        return {
            'self': st['self_nav'],
            'exalted': st['exalt_nav'],
            'debilitated': st['debil_nav'],
            'vargottama': st['vargottama'],
            'combust': False,
        }
    # default: rasi
    return {
        'self': st['self_rasi'],
        'exalted': st['exalt_rasi'],
        'debilitated': st['debil_rasi'],
        'vargottama': st['vargottama'],
        'combust': st['combust'],
    }

def fmt_planet_label(code, flags):
    base = HN_ABBR.get(code, code)
    if flags.get('exalted'): base += '↑'
    if flags.get('debilitated'): base += '↓'
    if flags.get('combust'): base += '^'
    return base



def planet_navamsa_house(lon_sid, nav_lagna_sign):
    # Return 1..12 house index in Navamsa for a planet
    nav_sign = navamsa_sign_from_lon_sid(lon_sid)  # 1..12
    return ((nav_sign - nav_lagna_sign) % 12) + 1

def build_navamsa_house_planets(sidelons, nav_lagna_sign):
    # Map: house -> list of planet abbreviations in Navamsa
    house_map = {i: [] for i in range(1, 13)}
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        h = planet_navamsa_house(sidelons[code], nav_lagna_sign)
        house_map[h].append(HN_ABBR.get(code, code))
    return house_map


def build_rasi_house_planets_marked(sidelons, lagna_sign):
    house_map = {i: [] for i in range(1, 13)}
    stats = compute_statuses_all(sidelons)
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        sign = planet_rasi_sign(sidelons[code])
        h = ((sign - lagna_sign) % 12) + 1
        fl = _make_flags('rasi', stats[code])
        label = fmt_planet_label(code, fl)
        house_map[h].append({'txt': label, 'flags': fl})
    return house_map

def build_navamsa_house_planets_marked(sidelons, nav_lagna_sign):
    house_map = {i: [] for i in range(1, 13)}
    stats = compute_statuses_all(sidelons)
    sun_nav = stats['Su']['nav']  # Sun's Navāṁśa sign
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        nav_sign = navamsa_sign_from_lon_sid(sidelons[code])
        h = ((nav_sign - nav_lagna_sign) % 12) + 1
        fl = _make_flags('nav', stats[code])   # nav-based self/exalt/debil
        # Navāṁśa combust rule: planet combust iff shares Nav sign with Sun
        if code not in ('Su','Ra','Ke'):
            fl['combust'] = (nav_sign == sun_nav)
        else:
            fl['combust'] = False
        label = fmt_planet_label(code, fl)
        house_map[h].append({'txt': label, 'flags': fl})
    return house_map


def build_rasi_house_planets(sidelons, lagna_sign):
    # Map: house -> list of planet abbreviations in Rasi (Lagna) chart
    house_map = {i: [] for i in range(1, 13)}
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        sign = int(sidelons[code] // 30) + 1  # 1..12
        h = ((sign - lagna_sign) % 12) + 1
        house_map[h].append(HN_ABBR.get(code, code))
    return house_map

def _apply_hindi_caption_style(paragraph, size_pt=11, underline=True, bold=True):
    if not paragraph.runs:
        paragraph.add_run("")
    r = paragraph.runs[0]
    r.bold = bold; r.underline = underline; r.font.size = Pt(size_pt)
    rpr = r._element.rPr or OxmlElement('w:rPr')
    if r._element.rPr is None: r._element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
    if rpr.find(qn('w:rFonts')) is None: rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), HINDI_FONT)

def set_sidereal_locked():
    swe.set_sid_mode(AYANAMSHA_VAL, 0, 0)

def dms_exact(deg):
    d = int(deg); m_float = (deg - d) * 60.0; m = int(m_float); s = (m_float - m) * 60.0
    return d, m, s

def fmt_deg_sign(lon_sid):
    sign=int(lon_sid//30) + 1; deg_in_sign = lon_sid % 30.0
    d,m,s=dms_exact(deg_in_sign); s_rounded = int(round(s))
    if s_rounded == 60: s_rounded = 0; m += 1
    if m == 60: m = 0; d += 1; 
    if d == 30: d = 0
    return sign, f"{d:02d}°{m:02d}'{s_rounded:02d}\""

def kp_sublord(lon_sid):
    NAK=360.0/27.0
    ORDER = ['Ke','Ve','Su','Mo','Ma','Ra','Ju','Sa','Me']
    YEARS = {'Ke':7,'Ve':20,'Su':6,'Mo':10,'Ma':7,'Ra':18,'Ju':16,'Sa':19,'Me':17}
    part = lon_sid % 360.0; ni = int(part // NAK); pos = part - ni*NAK
    lord = ORDER[ni % 9]; start = ORDER.index(lord)
    seq = [ORDER[(start+i)%9] for i in range(9)]
    acc = 0.0
    for L in seq:
        seg = NAK * (YEARS[L]/120.0)
        if pos <= acc + seg + 1e-9: return lord, L
        acc += seg
    return lord, seq[-1]

def geocode(place, api_key):
    if not api_key: raise RuntimeError("Geoapify key missing. Add GEOAPIFY_API_KEY in Secrets.")
    base="https://api.geoapify.com/v1/geocode/search?"
    q = urllib.parse.urlencode({"text":place, "format":"json", "limit":1, "apiKey":api_key})
    with urllib.request.urlopen(base+q, timeout=15) as r: j = json.loads(r.read().decode())
    if j.get("results"):
        res=j["results"][0]; return float(res["lat"]), float(res["lon"]), res.get("formatted", place)
    raise RuntimeError("Place not found.")


def get_timezone_offset_simple(lat, lon):
    """Simple timezone offset calculation for auto-population using hardcoded values"""
    try:
        tf = TimezoneFinder()
        tzname = tf.timezone_at(lat=lat, lng=lon)
        
        # Hardcoded timezone offsets to avoid pytz issues
        timezone_offsets = {
            'Asia/Kolkata': 5.5,           # India
            'Asia/Dubai': 4.0,             # UAE
            'Asia/Karachi': 5.0,           # Pakistan
            'Asia/Dhaka': 6.0,             # Bangladesh
            'Asia/Kathmandu': 5.75,        # Nepal
            'Asia/Colombo': 5.5,           # Sri Lanka
            'America/New_York': -5.0,      # EST (US East)
            'America/Chicago': -6.0,       # CST (US Central)
            'America/Denver': -7.0,        # MST (US Mountain)
            'America/Los_Angeles': -8.0,   # PST (US West)
            'Europe/London': 0.0,          # GMT (UK)
            'Europe/Paris': 1.0,           # CET (France, Germany, etc.)
            'Europe/Moscow': 3.0,          # MSK (Russia)
            'Asia/Tokyo': 9.0,             # JST (Japan)
            'Asia/Shanghai': 8.0,          # CST (China)
            'Australia/Sydney': 10.0,      # AEST (Australia East)
            'Australia/Perth': 8.0,        # AWST (Australia West)
            'Africa/Johannesburg': 2.0,    # SAST (South Africa)
            'America/Sao_Paulo': -3.0,     # BRT (Brazil)
            'America/Mexico_City': -6.0,   # CST (Mexico)
            'America/Toronto': -5.0,       # EST (Canada)
        }
        
        if tzname in timezone_offsets:
            offset = timezone_offsets[tzname]
            return offset
        else:
            print(f"DEBUG: Unknown timezone {tzname}, defaulting to 0.0")
            return 0.0
            
    except Exception as e:
        print(f"DEBUG: Timezone detection failed: {e}")
        return 0.0

def tz_from_latlon(lat, lon, dt_local):
    tf = TimezoneFinder()
    tzname = tf.timezone_at(lat=lat, lng=lon)
    
    # Debug output for timezone detection
    print(f"DEBUG: Coordinates: lat={lat}, lon={lon}")
    print(f"DEBUG: TimezoneFinder detected: {tzname}")
    
    # Fallback if no timezone detected by TimezoneFinder
    if not tzname:
        tzname = "Etc/UTC"
        print(f"DEBUG: No timezone detected by TimezoneFinder, falling back to UTC")
    
    try:
        # Create a fresh naive datetime to avoid any timezone issues
        clean_dt = datetime.datetime(dt_local.year, dt_local.month, dt_local.day, 
                                   dt_local.hour, dt_local.minute, dt_local.second)
        
        tz = pytz.timezone(tzname)
        dt_local_aware = tz.localize(clean_dt)
        dt_utc_naive = dt_local_aware.astimezone(pytz.utc).replace(tzinfo=None)
        offset_hours = tz.utcoffset(dt_local_aware).total_seconds()/3600.0
        print(f"DEBUG: Successfully processed timezone: {tzname}, offset: {offset_hours} hours")
        return tzname, offset_hours, dt_utc_naive
    except Exception as e:
        print(f"DEBUG: Timezone processing error: {e}")
        # For auto-population, we just need the offset, so let's calculate it directly
        try:
            tz = pytz.timezone(tzname)
            # Use a simple reference datetime to get the offset
            ref_dt = datetime.datetime(2025, 6, 15, 12, 0, 0)  # Mid-year to avoid DST issues
            ref_aware = tz.localize(ref_dt)
            offset_hours = tz.utcoffset(ref_aware).total_seconds()/3600.0
            print(f"DEBUG: Direct offset calculation successful: {offset_hours} hours")
            return tzname, offset_hours, dt_local.replace(tzinfo=None) if hasattr(dt_local, 'tzinfo') else dt_local
        except Exception as e2:
            print(f"DEBUG: Direct offset calculation also failed: {e2}, falling back to UTC")
            return "Etc/UTC", 0.0, dt_local.replace(tzinfo=None) if hasattr(dt_local, 'tzinfo') else dt_local


def sidereal_positions(dt_utc):
    jd = swe.julday(dt_utc.year, dt_utc.month, dt_utc.day, dt_utc.hour + dt_utc.minute/60 + dt_utc.second/3600)
    set_sidereal_locked(); flags = swe.FLG_SWIEPH | swe.FLG_SPEED | swe.FLG_SIDEREAL
    out = {}
    for code, p in [('Su',swe.SUN),('Mo',swe.MOON),('Ma',swe.MARS),('Me',swe.MERCURY),('Ju',swe.JUPITER),('Ve',swe.VENUS),('Sa',swe.SATURN)]:
        xx,_ = swe.calc_ut(jd, p, flags); out[code] = xx[0] % 360.0
    xx,_ = swe.calc_ut(jd, swe.MEAN_NODE, flags)  # Mean node locked
    out['Ra'] = xx[0] % 360.0; out['Ke'] = (out['Ra'] + 180.0) % 360.0
    ay = swe.get_ayanamsa_ut(jd); return jd, ay, out

def ascendant_sign(jd, lat, lon, ay):
    cusps, ascmc = swe.houses_ex(jd, lat, lon, b'P'); asc_trop = ascmc[0]; asc_sid = (asc_trop - ay) % 360.0
    return int(asc_sid // 30) + 1, asc_sid

def navamsa_sign_from_lon_sid(lon_sid):
    sign = int(lon_sid // 30) + 1; deg_in_sign = lon_sid % 30.0; pada = int(deg_in_sign // (30.0/9.0))
    if sign in (1,4,7,10): start = sign
    elif sign in (2,5,8,11): start = ((sign + 8 - 1) % 12) + 1
    else: start = ((sign + 4 - 1) % 12) + 1
    return ((start - 1 + pada) % 12) + 1

def positions_table_no_symbol(sidelons):
    rows=[]
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        lon=sidelons[code]; sign, deg_str = fmt_deg_sign(lon); nak_lord, sub_lord = kp_sublord(lon)
        rows.append([HN[code], sign, deg_str, HN[nak_lord], HN[sub_lord]])
    return pd.DataFrame(rows, columns=["ग्रह","राशि","अंश","नक्षत्र","उप‑नक्षत्र"])

ORDER = ['Ke','Ve','Su','Mo','Ma','Ra','Ju','Sa','Me']
YEARS = {'Ke':7,'Ve':20,'Su':6,'Mo':10,'Ma':7,'Ra':18,'Ju':16,'Sa':19,'Me':17}

def moon_balance_days(moon_sid):
    NAK=360.0/27.0; part = moon_sid % 360.0; ni = int(part // NAK); pos = part - ni*NAK
    md_lord = ORDER[ni % 9]; frac = pos/NAK; remaining_days = YEARS[md_lord]*(1 - frac)*YEAR_DAYS
    return md_lord, remaining_days

def build_mahadashas_days_utc(birth_utc_dt, moon_sid):
    md_lord, rem_days = moon_balance_days(moon_sid); end_limit = birth_utc_dt + datetime.timedelta(days=100*YEAR_DAYS)
    segments=[]; birth_md_start = birth_utc_dt; birth_md_end = min(birth_md_start + datetime.timedelta(days=rem_days), end_limit)
    segments.append({"planet": md_lord, "start": birth_md_start, "end": birth_md_end, "days": rem_days})
    idx = (ORDER.index(md_lord) + 1) % 9; t = birth_md_end
    while t < end_limit:
        L = ORDER[idx]; dur_days = YEARS[L]*YEAR_DAYS; end = min(t + datetime.timedelta(days=dur_days), end_limit)
        segments.append({"planet": L, "start": t, "end": end, "days": dur_days}); t = end; idx = (idx + 1) % 9
    return segments
# --- FIXED: compact kundali rendering with zero padding ---
def render_north_diamond(size_px=800, stroke=3):
    fig, ax = plt.subplots(figsize=(size_px/200, size_px/200), dpi=200)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.set_aspect('equal')
    ax.axis('off')
    # Outer square
    ax.plot([0,1,1,0,0],[0,0,1,1,0], linewidth=stroke, color='black')
    # Diagonals
    ax.plot([0,1],[1,0], linewidth=stroke, color='black')
    ax.plot([0,1],[0,1], linewidth=stroke, color='black')
    # Midpoint diamond
    ax.plot([0,0.5],[0.5,1], linewidth=stroke, color='black')
    ax.plot([0.5,1],[1,0.5], linewidth=stroke, color='black')
    ax.plot([1,0.5],[0.5,0], linewidth=stroke, color='black')
    ax.plot([0.5,0],[0,0.5], linewidth=stroke, color='black')
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', pad_inches=0)  # zero padding
    plt.close(fig); buf.seek(0); return buf

def rotated_house_labels(lagna_sign):
    order = [str(((lagna_sign - 1 + i) % 12) + 1) for i in range(12)]
    return {"1":order[0],"2":order[1],"3":order[2],"4":order[3],"5":order[4],"6":order[5],"7":order[6],"8":order[7],"9":order[8],"10":order[9],"11":order[10],"12":order[11]}


def kundali_with_planets(size_pt=None, lagna_sign=1, house_planets=None):
    
    # robust default for size_pt so definition never depends on globals
    if size_pt is None:
        try:
            size_pt = CHART_W_PT
        except Exception:
            size_pt = 318  # safe fallback
# Like kundali_w_p_with_centroid_labels but adds small side-by-side planet boxes below the number
    if house_planets is None:
        house_planets = {i: [] for i in range(1, 13)}
    S=size_pt; L,T,R,B=0,0,S,S
    TM=(S/2,0); RM=(S,S/2); BM=(S/2,S); LM=(0,S/2)
    P_lt=(S/4,S/4); P_rt=(3*S/4,S/4); P_rb=(3*S/4,3*S/4); P_lb=(S/4,3*S/4); O=(S/2,S/2)
    labels = rotated_house_labels(lagna_sign)
    houses = {
        "1":[TM,P_rt,O,P_lt],
        "2":[(0,0),TM,P_lt],
        "3":[(0,0),LM,P_lt],
        "4":[LM,O,P_lt,P_lb],
        "5":[LM,(0,S),P_lb],
        "6":[(0,S),BM,P_lb],
        "7":[BM,P_rb,O,P_lb],
        "8":[BM,(S,S),P_rb],
        "9":[RM,(S,S),P_rb],
        "10":[RM,O,P_rt,P_rb],
        "11":[(S,0),RM,P_rt],
        "12":[TM,(S,0),P_rt],
    }
    def centroid(poly):
        A=Cx=Cy=0.0; n=len(poly)
        for i in range(n):
            x1,y1=poly[i]; x2,y2=poly[(i+1)%n]
            cross=x1*y2 - x2*y1
            A+=cross; Cx+=(x1+x2)*cross; Cy+=(y1+y2)*cross
        A*=0.5
        if abs(A)<1e-9:
            xs,ys=zip(*poly); return (sum(xs)/n, sum(ys)/n)
        return (Cx/(6*A), Cy/(6*A))
    # Style for house-number boxes
    style = HOUSE_NUM_STYLE.lower()
    if style == 'plain':
        NUM_FILL, NUM_STROKE, NUM_STROKE_W = '#ffffff', 'none', '0pt'
    elif style == 'bordered':
        NUM_FILL, NUM_STROKE, NUM_STROKE_W = '#ffffff', 'black', f'{HOUSE_NUM_BORDER_PT}pt'
    elif style == 'shaded':
        NUM_FILL, NUM_STROKE, NUM_STROKE_W = HOUSE_NUM_SHADE, 'none', '0pt'
    else:  # bordered_shaded
        NUM_FILL, NUM_STROKE, NUM_STROKE_W = HOUSE_NUM_SHADE, 'black', f'{HOUSE_NUM_BORDER_PT}pt'
    num_boxes=[]; planet_boxes=[]; occupied_rects=[]
    num_w=NUM_W_PT; num_h=NUM_H_PT; p_w,p_h=PLANET_W_PT,PLANET_H_PT; gap_x=GAP_X_PT; offset_y=OFFSET_Y_PT
    for k,poly in houses.items():
        bbox = _bbox_of_poly(poly)
        # house number box
        x,y = centroid(poly); left = x - num_w/2; top = y - num_h/2; txt = labels[k]
        left, top = _clamp_in_bbox(left, top, num_w, num_h, bbox, pad=2)

        nl, nt = _nudge_number_box(left, top, num_w, num_h, S, occupied_rects)
        left, top = nl, nt
        occupied_rects.append({'left': left, 'top': top, 'right': left + num_w, 'bottom': top + num_h});
        num_boxes.append(f'''
        <v:rect style="position:absolute;left:{left}pt;top:{top}pt;width:{num_w}pt;height:{num_h}pt;z-index:80" fillcolor="#ffffff" strokecolor="none" strokeweight="0pt">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:t>{txt}</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        ''')
        # planet row below number
        planets = house_planets.get(int(k), [])
        if planets:
            n = len(planets)
            max_cols = 2  # wrap after this many per row
            rows = (n + max_cols - 1) // max_cols
            gap_y = 2
            # compute total grid height and top start
            total_h = rows * p_h + (rows - 1) * gap_y
            # start rows just below the number box
            grid_top = y + (p_h/2 + 2) + offset_y
            for idx, pl in enumerate(planets):
                # normalize input item
                if isinstance(pl, dict):
                    label = str(pl.get('txt', '')).strip() or '?'
                    fl = pl.get('flags', {}) or {}
                else:
                    label = str(pl).strip() or '?'
                    fl = {}
                r = idx // max_cols
                c = idx % max_cols
                # columns in this row (last row can be shorter)
                cols_this = max_cols if r < rows - 1 else (n - max_cols * (rows - 1)) or max_cols
                row_w = cols_this * p_w + (cols_this - 1) * gap_x
                row_left = x - row_w / 2
                top_box = grid_top + r * (p_h + gap_y) - p_h / 2
                # keep within chart square bounds with margin and tiny shrink on edges
                M = 5
                row_left = max(M, min(row_left, S - row_w - M))
                top_box  = max(M, min(top_box,  S - p_h - M))
                edge_touch = (row_left <= M + 0.05) or (row_left >= S - row_w - M - 0.05) or (top_box <= M + 0.05) or (top_box >= S - p_h - M - 0.05)
                pw = p_w - (1 if edge_touch else 0)
                ph = p_h - (1 if edge_touch else 0)
                left_pl = row_left + c * (pw + gap_x)
                box_xml = (
                    f"<v:rect style=\"position:absolute;left:{left_pl}pt;top:{top_box}pt;width:{pw}pt;height:{ph}pt;z-index:6\" strokecolor=\"none\">"
                    + "<v:textbox inset=\"0,0,0,0\">"
                    + "<w:txbxContent xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
                    + f"<w:p><w:pPr><w:jc w:val=\"center\"/></w:pPr><w:r><w:t>{_xml_text(label)}</w:t></w:r></w:p>"
                    + "</w:txbxContent>"
                    + "</v:textbox>"
                    + "</v:rect>"
                )
                planet_boxes.append(box_xml)
                # overlays
                try:
                    selfr = bool(fl.get('self'))
                    varg  = bool(fl.get('vargottama'))
                except Exception:
                    selfr = varg = False
                if selfr:
                    circle_left = left_pl + 2
                    circle_top  = top_box + 1
                    circle_w    = pw - 4
                    circle_h    = ph - 2
                    oval_xml = (
                        f"<v:oval style=\"position:absolute;left:{circle_left}pt;top:{circle_top}pt;width:{circle_w}pt;height:{circle_h}pt;z-index:7\" fillcolor=\"none\" strokecolor=\"black\" strokeweight=\"0.75pt\"/>"
                    )
                    planet_boxes.append(oval_xml)
                if varg:
                    badge_w = 5; badge_h = 5
                    badge_left = left_pl + pw - badge_w + 0.5
                    badge_top  = top_box - 2
                    badge_xml = (
                        f"<v:rect style=\"position:absolute;left:{badge_left}pt;top:{badge_top}pt;width:{badge_w}pt;height:{badge_h}pt;z-index:8\" fillcolor=\"#ffffff\" strokecolor=\"black\" strokeweight=\"0.75pt\"/>"
                    )
                    planet_boxes.append(badge_xml)
    # Compose shapes after processing all houses
    boxes_xml = "\\n".join(num_boxes + planet_boxes)

    xml = f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r>
      <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word"><w10:wrap type="topAndBottom"/>
        <v:group style="position:relative;margin-left:0;margin-top:0;width:{S}pt;height:{int(S*0.80)}pt" coordorigin="0,0" coordsize="{S},{S}">
          <v:rect style="position:absolute;left:0;top:0;width:{S}pt;height:{S}pt;z-index:1" strokecolor="#CC6600" strokeweight="3pt" fillcolor="#ffdcc8"/>
          <v:line style="position:absolute;z-index:2" from="{L},{T}" to="{R},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{R},{T}" to="{L},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{S/2},{T}" to="{R},{S/2}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{R},{S/2}" to="{S/2},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{S/2},{B}" to="{L},{S/2}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{L},{S/2}" to="{S/2},{T}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          {boxes_xml}
        </v:group>
      </w:pict>
    </w:r></w:p>
    '''
    return parse_xml(xml)



def kundali_single_box(size_pt=220, lagna_sign=1, house_planets=None):
    # One text box per house: first row = house number, second row = planets (centered)
    if house_planets is None:
        house_planets = {i: [] for i in range(1, 13)}
    S=size_pt; L,T,R,B=0,0,S,S
    TM=(S/2,0); RM=(S,S/2); BM=(S/2,S); LM=(0,S/2)
    P_lt=(S/4,S/4); P_rt=(3*S/4,S/4); P_rb=(3*S/4,3*S/4); P_lb=(S/4,3*S/4); O=(S/2,S/2)
    labels = rotated_house_labels(lagna_sign)
    houses = {
        "1":[TM,P_rt,O,P_lt],
        "2":[(0,0),TM,P_lt],
        "3":[(0,0),LM,P_lt],
        "4":[LM,O,P_lt,P_lb],
        "5":[LM,(0,S),P_lb],
        "6":[(0,S),BM,P_lb],
        "7":[BM,P_rb,O,P_lb],
        "8":[BM,(S,S),P_rb],
        "9":[RM,(S,S),P_rb],
        "10":[RM,O,P_rt,P_rb],
        "11":[(S,0),RM,P_rt],
        "12":[TM,(S,0),P_rt],
    }
    def centroid(poly):
        A=Cx=Cy=0.0; n=len(poly)
        for i in range(n):
            x1,y1=poly[i]; x2,y2=poly[(i+1)%n]
            cross=x1*y2 - x2*y1
            A+=cross; Cx+=(x1+x2)*cross; Cy+=(y1+y2)*cross
        A*=0.5
        if abs(A)<1e-9:
            xs,ys=zip(*poly); return (sum(xs)/n, sum(ys)/n)
        return (Cx/(6*A), Cy/(6*A))
    box_w, box_h = 30, 26  # slightly taller to hold two lines cleanly
    text_boxes=[]
    for k,poly in houses.items():
        x,y = centroid(poly)
        left = x - box_w/2; top = y - box_h/2
        num = labels[k]
        pls = house_planets.get(int(k), [])
        if pls:
            planets_text = " ".join(pls)
            content = f'<w:r><w:t>{num}</w:t></w:r><w:r/><w:br/><w:r><w:t>{planets_text}</w:t></w:r>'
        else:
            content = f'<w:r><w:t>{num}</w:t></w:r>'
        text_boxes.append(f'''
        <v:rect style="position:absolute;left:{left}pt;top:{top}pt;width:{box_w}pt;height:{box_h}pt;z-index:5" strokecolor="none">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr>{content}</w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        ''')
    boxes_xml = "\\n".join(text_boxes)
    xml = f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r>
      <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word"><w10:wrap type="topAndBottom"/>
        <v:group style="position:relative;margin-left:0;margin-top:0;width:{S}pt;height:{int(S*0.80)}pt" coordorigin="0,0" coordsize="{S},{S}">
          <v:rect style="position:absolute;left:0;top:0;width:{S}pt;height:{S}pt;z-index:1" strokecolor="#CC6600" strokeweight="3pt" fillcolor="#ffdcc8"/>
          <v:line style="position:absolute;z-index:2" from="{L},{T}" to="{R},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{R},{T}" to="{L},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{S/2},{T}" to="{R},{S/2}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{R},{S/2}" to="{S/2},{B}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{S/2},{B}" to="{L},{S/2}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          <v:line style="position:absolute;z-index:2" from="{L},{S/2}" to="{S/2},{T}" strokecolor="#CC6600" strokeweight="1.25pt"/>
          {boxes_xml}
        </v:group>
      </w:pict>
    </w:r></w:p>
    '''
    return parse_xml(xml)


def kundali_w_p_with_centroid_labels(size_pt=220, lagna_sign=1):
    S=size_pt; TM=(S/2,0); RM=(S,S/2); BM=(S/2,S); LM=(0,S/2); P_lt=(S/4,S/4); P_rt=(3*S/4,S/4); P_rb=(3*S/4,3*S/4); P_lb=(S/4,3*S/4); O=(S/2,S/2)
    labels = rotated_house_labels(lagna_sign)
    houses = {"1":[TM,P_rt,O,P_lt],"2":[(0,0),TM,P_lt],"3":[(0,0),LM,P_lt],"4":[LM,O,P_lt,P_lb],"5":[LM,(0,S),P_lb],"6":[(0,S),BM,P_lb],"7":[BM,P_rb,O,P_lb],"8":[BM,(S,S),P_rb],"9":[RM,(S,S),P_rb],"10":[RM,O,P_rt,P_rb],"11":[(S,0),RM,P_rt],"12":[TM,(S,0),P_rt]}
    def centroid(poly):
        A=Cx=Cy=0.0; n=len(poly)
        for i in range(n):
            x1,y1=poly[i]; x2,y2=poly[(i+1)%n]; cross=x1*y2 - x2*y1; A+=cross; Cx+=(x1+x2)*cross; Cy+=(y1+y2)*cross
        A*=0.5
        if abs(A)<1e-9: xs,ys=zip(*poly); return (sum(xs)/n, sum(ys)/n)
        return (Cx/(6*A), Cy/(6*A))
    w=h=20; boxes=[]
    for k,poly in houses.items():
        x,y = centroid(poly); left = x - w/2; top = y - h/2; txt = labels[k]
        boxes.append(f'''
        <v:rect style="position:absolute;left:{left}pt;top:{top}pt;width:{w}pt;height:{h}pt;z-index:5" strokecolor="none">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:t>{txt}</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>''')
    boxes_xml = "\\n".join(boxes)
    xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r>
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word"><w10:wrap type="topAndBottom"/>
          <v:group style="position:relative;margin-left:0;margin-top:0;width:{S}pt;height:{int(S*0.80)}pt" coordorigin="0,0" coordsize="{S},{S}">
            <v:rect style="position:absolute;left:0;top:0;width:{S}pt;height:{S}pt;z-index:1" strokecolor="black" strokeweight="1.25pt" fillcolor="#ffdcc8"/>
            <v:line style="position:absolute;z-index:2" from="0,0" to="{S},{S}" strokecolor="black" strokeweight="1.25pt"/>
            <v:line style="position:absolute;z-index:2" from="{S},0" to="0,{S}" strokecolor="black" strokeweight="1.25pt"/>
            <v:line style="position:absolute;z-index:2" from="{S/2},0" to="{S},{S/2}" strokecolor="black" strokeweight="1.25pt"/>
            <v:line style="position:absolute;z-index:2" from="{S},{S/2}" to="{S/2},{S}" strokecolor="black" strokeweight="1.25pt"/>
            <v:line style="position:absolute;z-index:2" from="{S/2},{S}" to="0,{S/2}" strokecolor="black" strokeweight="1.25pt"/>
            <v:line style="position:absolute;z-index:2" from="0,{S/2}" to="{S/2},0" strokecolor="black" strokeweight="1.25pt"/>
            {boxes_xml}
          </v:group>
        </w:pict></w:r></w:p>'''
    return parse_xml(xml)

def add_table_borders(table, size=6):
    tbl = table._tbl; tblPr = tbl.tblPr; tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(size)); tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_font(table, pt=8.0):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(pt)

def center_header_row(table):
    for cell in table.rows[0].cells:
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if par.runs: par.runs[0].bold = True

# ===== MODERN DESIGN STYLING FUNCTIONS =====

def create_cylindrical_section_header(container, title_text, width_pt=320):
    """Create modern cylindrical tube-shaped section headers with dynamic width"""
    # Create paragraph for the header
    header_para = container.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    
    # Add the title text with styling
    run = header_para.add_run(title_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Add beautiful gradient background styling using VML shape 
    xml_content = f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:pPr>
        <w:jc w:val="center"/>
        <w:spacing w:before="120" w:after="100"/>
      </w:pPr>
      <w:r>
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word"><w10:wrap type="topAndBottom"/>
          <v:roundrect style="position:relative;width:{width_pt}pt;height:28pt;margin-left:auto;margin-right:auto" 
                       arcsize="45%" strokecolor="#D2691E" strokeweight="1.5pt">
            <v:fill type="gradient" color="#F15A23" color2="#FFEACC" angle="90" opacity="1"/>
            <v:textbox inset="8pt,4pt,8pt,4pt">
              <w:txbxContent>
                <w:p>
                  <w:pPr><w:jc w:val="center"/></w:pPr>
                  <w:r>
                    <w:rPr>
                      <w:color w:val="FFFFFF"/>
                      <w:sz w:val="24"/>
                      <w:b/>
                      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    </w:rPr>
                    <w:t>{title_text}</w:t>
                  </w:r>
                </w:p>
              </w:txbxContent>
            </v:textbox>
          </v:roundrect>
        </w:pict>
      </w:r>
    </w:p>'''
    
    try:
        from docx.oxml import parse_xml
        header_element = parse_xml(xml_content)
        container._element.append(header_element)
        # Remove the original paragraph we added
        container._element.remove(header_para._element)
    except Exception:
        # Fallback to simple styled text if VML fails
        pass
    # Ensure spacing after header so following table starts below the bar
    try:
        spacer = container.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
    except Exception:
        pass

def create_unified_personal_details_box(container, name, dob, tob, place):
    """Create single rounded corner box with title inside, matching reference image exactly"""
    
    # Try to create a rounded rectangle using VML for truly rounded corners
    try:
        # Create content text first
        content_text = f'''व्यक्तिगत विवरण

नाम: {name}
जन्म तिथि: {dob}
जन्म समय: {tob}
स्थान: {place}'''
        
        # Create VML rounded rectangle
        xml_content = f'''
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:pPr>
            <w:spacing w:before="0" w:after="120"/>
          </w:pPr>
          <w:r>
            <w:pict xmlns:v="urn:schemas-microsoft-com:vml">
              <v:roundrect style="position:relative;width:332pt;height:130pt" 
                           arcsize="15%" fillcolor="white" strokecolor="#F15A23" strokeweight="1.5pt">
                <v:textbox inset="12pt,10pt,12pt,10pt">
                  <w:txbxContent>
                    <w:p>
                      <w:pPr><w:jc w:val="center"/><w:spacing w:after="120"/></w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:color w:val="F15A23"/>
                          <w:sz w:val="22"/>
                          <w:b/>
                          <w:u/>
                        </w:rPr>
                        <w:t>व्यक्तिगत विवरण</w:t>
                      </w:r>
                    </w:p>
                    <w:p>
                      <w:pPr>
                        <w:spacing w:after="80"/>
                        <w:tabs>
                          <w:tab w:val="left" w:pos="1440"/>
                        </w:tabs>
                      </w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:color w:val="F15A23"/>
                          <w:sz w:val="20"/>
                          <w:b/>
                          <w:u/>
                        </w:rPr>
                        <w:t>नाम :</w:t>
                      </w:r>
                      <w:r>
                        <w:tab/>
                        <w:rPr>
                          <w:color w:val="000000"/>
                          <w:sz w:val="20"/>
                        </w:rPr>
                        <w:t>{name}</w:t>
                      </w:r>
                    </w:p>
                    <w:p>
                      <w:pPr>
                        <w:spacing w:after="80"/>
                        <w:tabs>
                          <w:tab w:val="left" w:pos="1440"/>
                        </w:tabs>
                      </w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:color w:val="F15A23"/>
                          <w:sz w:val="20"/>
                          <w:b/>
                          <w:u/>
                        </w:rPr>
                        <w:t>जन्म तिथि :</w:t>
                      </w:r>
                      <w:r>
                        <w:tab/>
                        <w:rPr>
                          <w:color w:val="000000"/>
                          <w:sz w:val="20"/>
                        </w:rPr>
                        <w:t>{dob}</w:t>
                      </w:r>
                    </w:p>
                    <w:p>
                      <w:pPr>
                        <w:spacing w:after="80"/>
                        <w:tabs>
                          <w:tab w:val="left" w:pos="1440"/>
                        </w:tabs>
                      </w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:color w:val="F15A23"/>
                          <w:sz w:val="20"/>
                          <w:b/>
                          <w:u/>
                        </w:rPr>
                        <w:t>जन्म समय :</w:t>
                      </w:r>
                      <w:r>
                        <w:tab/>
                        <w:rPr>
                          <w:color w:val="000000"/>
                          <w:sz w:val="20"/>
                        </w:rPr>
                        <w:t>{tob}</w:t>
                      </w:r>
                    </w:p>
                    <w:p>
                      <w:pPr>
                        <w:spacing w:after="40"/>
                        <w:tabs>
                          <w:tab w:val="left" w:pos="1440"/>
                        </w:tabs>
                      </w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:color w:val="F15A23"/>
                          <w:sz w:val="20"/>
                          <w:b/>
                          <w:u/>
                        </w:rPr>
                        <w:t>स्थान :</w:t>
                      </w:r>
                      <w:r>
                        <w:tab/>
                        <w:rPr>
                          <w:color w:val="000000"/>
                          <w:sz w:val="20"/>
                        </w:rPr>
                        <w:t>{place}</w:t>
                      </w:r>
                    </w:p>
                  </w:txbxContent>
                </v:textbox>
              </v:roundrect>
            </w:pict>
          </w:r>
        </w:p>'''
        
        from docx.oxml import parse_xml
        rounded_element = parse_xml(xml_content)
        container._element.append(rounded_element)
        return None  # No table to return
        
    except Exception:
        # Fallback to table approach if VML fails
        pass
    
    # Fallback: Create a table with rounded corners for unified personal details
    detail_table = container.add_table(rows=1, cols=1)
    detail_table.autofit = False
    detail_table.columns[0].width = Inches(3.5)
    
    cell = detail_table.rows[0].cells[0]
    
    # Add Title "व्यक्तिगत विवरण" inside the box at the top - compact spacing
    title_para = cell.add_paragraph('व्यक्तिगत विवरण')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(11)  # Slightly smaller for compact
    title_run.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    title_para.paragraph_format.space_after = Pt(4)  # Reduced from 8
    title_para.paragraph_format.space_before = Pt(0)  # Reduced from 2
    
    # Add Name - compact spacing
    name_para = cell.add_paragraph()
    name_title = name_para.add_run('नाम: ')
    name_title.bold = True
    name_title.font.size = Pt(9)  # Smaller font for compact
    name_title.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    name_content = name_para.add_run(str(name))
    name_content.font.size = Pt(9)  # Smaller font for compact
    name_content.font.color.rgb = RGBColor(0, 0, 0)  # Black color like in reference
    name_para.paragraph_format.space_after = Pt(1)  # Reduced from 3
    
    # Add Date of Birth - compact spacing
    dob_para = cell.add_paragraph()
    dob_title = dob_para.add_run('जन्म तिथि: ')
    dob_title.bold = True
    dob_title.font.size = Pt(9)  # Smaller font for compact
    dob_title.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    dob_content = dob_para.add_run(str(dob))
    dob_content.font.size = Pt(9)  # Smaller font for compact
    dob_content.font.color.rgb = RGBColor(0, 0, 0)  # Black color like in reference
    dob_para.paragraph_format.space_after = Pt(1)  # Reduced from 3
    
    # Add Time of Birth - compact spacing
    tob_para = cell.add_paragraph()
    tob_title = tob_para.add_run('जन्म समय: ')
    tob_title.bold = True
    tob_title.font.size = Pt(9)  # Smaller font for compact
    tob_title.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    tob_content = tob_para.add_run(str(tob))
    tob_content.font.size = Pt(9)  # Smaller font for compact
    tob_content.font.color.rgb = RGBColor(0, 0, 0)  # Black color like in reference
    tob_para.paragraph_format.space_after = Pt(1)  # Reduced from 3
    
    # Add Place - compact spacing
    place_para = cell.add_paragraph()
    place_title = place_para.add_run('स्थान: ')
    place_title.bold = True
    place_title.font.size = Pt(9)  # Smaller font for compact
    place_title.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    place_content = place_para.add_run(str(place))
    place_content.font.size = Pt(9)  # Smaller font for compact
    place_content.font.color.rgb = RGBColor(0, 0, 0)  # Black color like in reference
    place_para.paragraph_format.space_after = Pt(0)  # Reduced from 2
    
    # Apply compact rounded corner styling with minimal padding
    try:
        cell_elem = cell._tc
        tcPr = cell_elem.get_or_add_tcPr()
        
        # Add rounded corner borders using dotted style for rounded appearance
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')  # Thin border
            border.set(qn('w:color'), 'F15A23')  # Orange color matching reference
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        # Minimal padding for compact 1-page format
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'left', 'bottom', 'right'):
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '80')  # Minimal padding for compact layout
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # Clean white background
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'FFFFFF')  # Pure white background
        tcPr.append(shd)
        
        # Add rounded corner effect using XML for better circular appearance
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '0')
        tcW.set(qn('w:type'), 'auto')
        tcPr.append(tcW)
        
    except Exception:
        pass
    
    return detail_table

def create_rounded_detail_box(container, title, content):
    """Create rounded corner boxes for personal details"""
    # Create a table with rounded corners for the detail box
    detail_table = container.add_table(rows=1, cols=1)
    detail_table.autofit = False
    detail_table.columns[0].width = Inches(6.0)
    
    cell = detail_table.rows[0].cells[0]
    
    # Add title
    title_para = cell.add_paragraph(title)
    title_run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(241, 90, 35)  # Orange color
    title_para.paragraph_format.space_after = Pt(2)
    
    # Add content
    content_para = cell.add_paragraph(content)
    content_run = content_para.runs[0] if content_para.runs else content_para.add_run(content)
    content_run.font.size = Pt(9)
    content_run.font.color.rgb = RGBColor(51, 51, 51)  # Dark grey
    
    # Apply rounded corner styling to the cell
    try:
        cell_elem = cell._tc
        tcPr = cell_elem.get_or_add_tcPr()
        
        # Add rounded corner borders
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:color'), 'F15A23')  # Dark orange
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        # Add cell padding
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'left', 'bottom', 'right'):
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '100')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
    except Exception:
        pass
    
    return detail_table

def create_rounded_table_container(doc, table_content, width_pt=400, height_pt=200):
    """Create a VML rounded rectangle container for tables with true circular corners"""
    # Create paragraph with VML roundrect container
    p = doc.add_paragraph()
    
    # Create VML roundrect with genuine rounded corners
    xml_content = f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:pPr>
        <w:spacing w:before="60" w:after="60"/>
      </w:pPr>
      <w:r>
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml">
          <v:roundrect style="position:relative;width:{width_pt}pt;height:{height_pt}pt" 
                       arcsize="15%" fillcolor="#ffdcc8" strokecolor="#D2691E" strokeweight="2pt">
            <v:textbox inset="8pt,8pt,8pt,8pt">
              <w:txbxContent>
                {table_content}
              </w:txbxContent>
            </v:textbox>
          </v:roundrect>
        </w:pict>
      </w:r>
    </w:p>
    '''
    
    return parse_xml(xml_content)

def apply_premium_table_style(table, header_color_rgb=(204, 102, 0), alt_row_color_rgb=(255, 235, 224)):
    """Apply premium professional styling to tables with genuine rounded corners using VML background"""
    try:
        # Apply table borders - no outer borders for rounded effect
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        
        # Apply rounded corner border styling
        border_styles = {
            'top': ('thick', '12'),     # Thick top border for rounded effect
            'left': ('thick', '12'),    # Thick left border for rounded effect 
            'bottom': ('thick', '12'),  # Thick bottom border for rounded effect
            'right': ('thick', '12'),   # Thick right border for rounded effect
            'insideH': ('single', '6'),  # Internal horizontal borders
            'insideV': ('single', '6')   # Internal vertical borders
        }
        
        for edge, (style, size) in border_styles.items():
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), style)
            border.set(qn('w:sz'), size)
            border.set(qn('w:color'), 'D2691E')  # Dark orange color
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Add table alignment
        tblAlign = OxmlElement('w:jc')
        tblAlign.set(qn('w:val'), 'center')
        tblPr.append(tblAlign)
        
        # Add table style properties for rounded corners
        try:
            # Apply table-level styling for rounded appearance
            tblStyle = OxmlElement('w:tblStyle')
            tblStyle.set(qn('w:val'), 'TableGrid')  # Use a style that supports rounding
            tblPr.insert(0, tblStyle)
            
            # Add table cell margins for better spacing
            tblCellMar = OxmlElement('w:tblCellMar')
            for side in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), '60')  # Add some margin
                margin.set(qn('w:type'), 'dxa')
                tblCellMar.append(margin)
            tblPr.append(tblCellMar)
            
        except Exception:
            pass
        
        # Add genuine VML rounded corners to corner cells
        try:
            # Get corner cells and add VML rounded rectangle backgrounds
            num_rows = len(table.rows)
            num_cols = len(table.rows[0].cells) if table.rows else 0
            
            if num_rows > 0 and num_cols > 0:
                # Apply VML rounded backgrounds to corner cells
                corner_positions = [
                    (0, 0, 'top-left'),
                    (0, num_cols-1, 'top-right'),
                    (num_rows-1, 0, 'bottom-left'),
                    (num_rows-1, num_cols-1, 'bottom-right')
                ]
                
                for row_idx, col_idx, corner_type in corner_positions:
                    try:
                        cell = table.cell(row_idx, col_idx)
                        
                        # Add VML rounded rectangle as paragraph inside the cell
                        vml_para = cell.add_paragraph()
                        
                        # Determine corner-specific arcsize
                        arcsize_map = {
                            'top-left': '0 0 20% 20%',
                            'top-right': '20% 0 0 20%', 
                            'bottom-left': '0 20% 20% 0',
                            'bottom-right': '20% 20% 0 0'
                        }
                        
                        # Create VML rounded corner element
                        vml_xml = f'''
                        <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                          <v:roundrect style="position:absolute;left:0;top:0;width:100%;height:100%;z-index:-1" 
                                       arcsize="15%" fillcolor="#ffdcc8" strokecolor="#D2691E" strokeweight="1pt">
                          </v:roundrect>
                        </w:pict>
                        '''
                        
                        # Parse and insert VML into the paragraph
                        vml_element = parse_xml(vml_xml)
                        vml_para._p.append(vml_element._element)
                        
                    except Exception:
                        continue
                        
        except Exception:
            pass
        
        # Style header row with premium look
        header_cells = table.rows[0].cells
        for cell in header_cells:
            # Premium header background
            cell_elem = cell._tc
            tcPr = cell_elem.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '{:02x}{:02x}{:02x}'.format(*header_color_rgb))
            tcPr.append(shd)
            
            # Add minimal cell padding for compactness
            tcMar = OxmlElement('w:tcMar')
            for side in ('top', 'left', 'bottom', 'right'):
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), '40')  # Reduced from 100 to 40
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            
            # Enhanced header text styling
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)  # Slightly smaller for compactness
                    run.font.name = 'Calibri'
        
        # Style data rows with professional alternating colors
        for i, row in enumerate(table.rows[1:], 1):
            for cell in row.cells:
                cell_elem = cell._tc
                tcPr = cell_elem.get_or_add_tcPr()
                
                # Alternating row colors: odd rows (1,3,5...) get beautiful light orange background
                if i % 2 == 1:  # Odd rows get the beautiful light orange background
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '{:02x}{:02x}{:02x}'.format(*alt_row_color_rgb))
                    tcPr.append(shd)
                # Even rows (2,4,6...) get no background color (default white)
                
                # Add minimal cell padding for all data cells
                tcMar = OxmlElement('w:tcMar')
                for side in ('top', 'left', 'bottom', 'right'):
                    margin = OxmlElement(f'w:{side}')
                    margin.set(qn('w:w'), '30')  # Reduced from 80 to 30
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                # Enhanced data text styling
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(51, 51, 51)
                        run.font.size = Pt(8)  # Even smaller for data cells to fit more content
                        run.font.name = 'Calibri'
    except Exception:
        pass

def create_section_header(container, title, color_rgb=(25, 55, 109)):
    """Create original decorative section header"""
    # Original section header styling
    h = container.add_paragraph(title)
    h.runs[0].bold = True
    h.runs[0].underline = True
    h.runs[0].font.size = Pt(13)
    h.runs[0].font.color.rgb = RGBColor(*color_rgb)
    h.runs[0].font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(6)
    
    return h

def set_col_widths(table, widths_inch):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_inch):
            row.cells[i].width = Inches(w)

def sanitize_filename(name: str) -> str:
    # Keep spaces; strip leading/trailing; allow letters/digits/space/_/- only
    raw = (name or 'Horoscope').strip()
    cleaned = ''.join(ch for ch in raw if ch.isalnum() or ch in ' _-')
    return cleaned or 'Horoscope'

def _utc_to_local(dt_utc, tzname, tz_hours, used_manual):
    if used_manual: return dt_utc + datetime.timedelta(hours=tz_hours)
    try:
        tz = pytz.timezone(tzname); return tz.fromutc(dt_utc.replace(tzinfo=pytz.utc))
    except Exception:
        return dt_utc + datetime.timedelta(hours=tz_hours)

# Core UI

def _house_from_lagna(sign:int, lagna_sign:int)->int:
    return ((sign - lagna_sign) % 12) + 1  # 1..12

def _english_bhav_label(h:int)->str:
    try:
        h_int = int(h)
    except Exception:
        return f"{h}वाँ भाव"
    return f"{h_int}वाँ भाव"

def detect_muntha_house(lagna_sign:int, dob_dt):
    # Approx: years elapsed since birth to today -> advance houses from lagna
    try:
        from datetime import datetime, timezone
        years = datetime.now(timezone.utc).year - dob_dt.year
        return ((lagna_sign - 1 + years) % 12) + 1
    except Exception:
        return None

def detect_sade_sati_or_dhaiyya(sidelons:dict, transit_dt=None):
    # Returns: (status, phase) where status in {"साढ़ेसाती", "शनि ढैय्या", None}
    # Uses *transit Saturn* vs *natal Moon*. Phase only if साढ़ेसाती: "प्रथम चरण" / "द्वितीय चरण" / "तृतीय चरण".
    try:
        # Natal Moon sign
        moon = planet_rasi_sign(sidelons['Mo'])
        # Transit Saturn sign at transit_dt (or now)
        from datetime import datetime, timezone
        if transit_dt is None:
            tdt = datetime.now(timezone.utc)
        else:
            tdt = transit_dt
        _jd, _ay, trans = sidereal_positions(tdt.replace(tzinfo=None) if hasattr(tdt, 'tzinfo') else tdt)
        sat = planet_rasi_sign(trans['Sa'])
        d = (sat - moon) % 12
        if d in (11, 0, 1):
            phase = {11: "प्रथम चरण", 0: "द्वितीय चरण", 1: "तृतीय चरण"}[d]
            return "साढ़ेसाती", phase
        if d in (3, 7):
            return "शनि ढैय्या", None
        return None, None
    except Exception:
        return None, None

def detect_kaalsarp(sidelons:dict)->bool:
    try:
        ra = sidelons['Ra'] % 360.0
        ke = (ra + 180.0) % 360.0
        span = (ke - ra) % 360.0  # should be 180
        inside = 0
        for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa']:
            ang = (sidelons[code] - ra) % 360.0
            if ang <= span:
                inside += 1
        return inside == 7
    except Exception:
        return False

def detect_chandal(sidelons:dict)->bool:
    try:
        ju = planet_rasi_sign(sidelons['Ju'])
        return ju == planet_rasi_sign(sidelons['Ra']) or ju == planet_rasi_sign(sidelons['Ke'])
    except Exception:
        return False

def detect_pitru(sidelons:dict)->bool:
    try:
        su = planet_rasi_sign(sidelons['Su'])
        return su == planet_rasi_sign(sidelons['Ra']) or su == planet_rasi_sign(sidelons['Ke'])
    except Exception:
        return False

def detect_neech_bhang(sidelons:dict, lagna_sign:int)->bool:
    try:
        stats = compute_statuses_all(sidelons)
        for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa']:
            if stats[code]['debil_rasi']:
                debil_sign = stats[code]['rasi']
                lord = SIGN_LORD.get(debil_sign)
                if lord and lord in sidelons:
                    lord_sign = planet_rasi_sign(sidelons[lord])
                    h = _house_from_lagna(lord_sign, lagna_sign)
                    if h in (1,4,7,10):
                        return True
        return False
    except Exception:
        return False

def compact_table_paragraphs(tbl):
    try:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
    except Exception:
        pass

def add_pramukh_bindu_section(container_cell, sidelons, lagna_sign, dob_dt):
    spacer = container_cell.add_paragraph("")
    spacer.paragraph_format.space_after = Pt(0)
    # Title
    # title = container_cell.add_paragraph("प्रमुख बिंदु")
    # # Match other section titles
    # _apply_hindi_caption_style(title, size_pt=11, underline=True, bold=True)
    # title.paragraph_format.space_before = Pt(0)
    # title.paragraph_format.space_after = Pt(2)
    # title.paragraph_format.space_before = Pt(6)
    # title.paragraph_format.space_after = Pt(3)
    create_cylindrical_section_header(container_cell, "प्रमुख बिंदु", width_pt=260)

    rows = []

    # Muntha
    m = detect_muntha_house(lagna_sign, dob_dt)
    if m:
        rows.append(("मुन्था (वर्तमान वर्ष)", _english_bhav_label(m)))

    # Sade Sati / Dhaiyya
    status, phase = detect_sade_sati_or_dhaiyya(sidelons)
    if status:
        rows.append(("साढ़ेसाती/शनि ढैय्या", status))
        if status == "साढ़ेसाती" and phase:
            rows.append(("साढ़ेसाती का चरण", phase))

    # Dosha/Yoga (only if True)
    if detect_kaalsarp(sidelons):
        rows.append(("कालसर्प दोष", "हाँ"))
    if detect_chandal(sidelons):
        rows.append(("चांडाल योग", "हाँ"))
    if detect_pitru(sidelons):
        rows.append(("पितृ दोष", "हाँ"))
    if detect_neech_bhang(sidelons, lagna_sign):
        rows.append(("नीच भंग राज योग", "हाँ"))

    if not rows:
        # Nothing to show; avoid adding an empty table
        return

    t = container_cell.add_table(rows=0, cols=2)
    t.autofit = True
    # Match font size with other tables
    try:
        set_table_font(t, pt=BASE_FONT_PT)
    except Exception:
        pass
    for left_txt, right_txt in rows:
        r = t.add_row().cells
        r[0].text = left_txt
        r[1].text = right_txt

    # Borders similar to other tables
    add_table_borders(t, size=6)
    apply_premium_table_style(t)  # Apply orange headers and alternating grey rows
    compact_table_paragraphs(t)
def main():
    pass
    # === Brand Header ===
    # === End Brand Header ===

    # st.title(APP_TITLE)  # removed to avoid duplicate brand name

# === CSS styling for input field sizes ===
st.markdown(
    """
    <style>
    /* Name and Place text input boxes - larger for full text visibility */
    div[data-testid="column"]:nth-child(1) .stTextInput > div > div > input,
    div[data-testid="column"]:nth-child(2) .stTextInput > div > div > input {
        width: 80% !important;
    }
    div[data-testid="column"]:nth-child(1) .stTextInput > div > div,
    div[data-testid="column"]:nth-child(2) .stTextInput > div > div {
        width: 80% !important;
    }
    /* All row 2 fields (DOB, TOB, UTC) - uniform professional sizing */
    .stDateInput > div > div > input,
    .stTimeInput > div > div > input {
        width: 80% !important;
    }
    .stDateInput > div > div,
    .stTimeInput > div > div {
        width: 80% !important;
    }
    /* Remove all UTC-specific sizing - let it stay natural for uniformity */
    /* Restore normal scrolling */
    /* Center align both buttons perfectly */
    div[data-testid="column"]:nth-child(2) .stDownloadButton > button {
        display: block !important;
        margin: 0 auto !important;
        margin-left: 25% !important;
        width: fit-content !important;
    }
    div[data-testid="column"]:nth-child(2) .stButton > button {
        display: block !important;
        margin: 0 auto !important;
        width: fit-content !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# === Set submitted state if button was clicked (needed for immediate validation) ===
if 'generate_clicked' not in st.session_state:
    st.session_state['generate_clicked'] = False

# === Reorganized form layout ===
# Row 1: Name and Place of Birth
row1c1, row1c2 = st.columns(2)
with row1c1:
    name_val = (st.session_state.get('name_input','') or '').strip()
    name_err = (st.session_state.get('submitted') or st.session_state.get('generate_clicked')) and (not name_val)
    render_label('Name <span style="color:red">*</span>', name_err)
    name = st.text_input("Name", key="name_input", label_visibility="collapsed")
with row1c2:
    place_val = (st.session_state.get('place_input','') or '').strip()
    place_err = (st.session_state.get('submitted') or st.session_state.get('generate_clicked')) and (not place_val)
    render_label('Place of Birth (City, State, Country) <span style="color:red">*</span>', place_err)
    place = st.text_input("Place of Birth", key="place_input", label_visibility="collapsed")

# Clear previous generation if any field changes
current_form_values = {
    'name': st.session_state.get('name_input', '').strip(),
    'place': st.session_state.get('place_input', '').strip(), 
    'dob': st.session_state.get('dob_input'),
    'tob': st.session_state.get('tob_input'),
    'tz': st.session_state.get('tz_input', '').strip()
}

last_form_values = st.session_state.get('last_form_values', {})

# Check if any field changed
form_changed = current_form_values != last_form_values
if form_changed and last_form_values:  # Don't clear on first load
    # Clear previous generation when any field changes
    st.session_state.pop('kundali_doc', None)
    st.session_state.pop('generation_completed', None)
    st.session_state.pop('submitted', None)

# Update last values
st.session_state['last_form_values'] = current_form_values

# Auto-populate UTC offset when place changes
place_input_val = st.session_state.get('place_input', '').strip()
if place_input_val and place_input_val != st.session_state.get('last_place_checked', ''):
    try:
        api_key = st.secrets.get("GEOAPIFY_API_KEY", "")
        if api_key:
            # Try to geocode and detect timezone
            lat, lon, disp = geocode(place_input_val, api_key)
            # Use simple timezone offset calculation for auto-population
            offset_hours = get_timezone_offset_simple(lat, lon)
            # Auto-populate the UTC offset field
            st.session_state['tz_input'] = str(offset_hours)
            st.session_state['last_place_checked'] = place_input_val
            st.rerun()  # Refresh to show the auto-populated value
    except Exception as e:
        # If auto-detection fails, just leave the field for manual entry
        pass

# Row 2: Date of Birth, Time of Birth, and UTC offset override
row2c1, row2c2, row2c3 = st.columns(3)
with row2c1:
    # Check validation using current session state (widget will update it)
    dob_current = st.session_state.get('dob_input', datetime.date.today())
    dob_err = (st.session_state.get('submitted') or st.session_state.get('generate_clicked')) and (dob_current is None)
    render_label('Date of Birth <span style="color:red">*</span>', dob_err)
    dob = st.date_input("Date of Birth", key="dob_input", label_visibility="collapsed",
                        min_value=datetime.date(1800,1,1), max_value=datetime.date(2100,12,31))
with row2c2:
    # Check validation using current session state (widget will update it)
    tob_current = st.session_state.get('tob_input', datetime.time(12, 0))
    tob_err = (st.session_state.get('submitted') or st.session_state.get('generate_clicked')) and (tob_current is None)
    render_label('Time of Birth <span style="color:red">*</span>', tob_err)
    tob = st.time_input("Time of Birth", key="tob_input", label_visibility="collapsed", step=datetime.timedelta(minutes=1))
with row2c3:
    tz_val = (st.session_state.get('tz_input','') or '').strip()
    place_val = (st.session_state.get('place_input','') or '').strip()
    tz_err = (st.session_state.get('submitted') or st.session_state.get('generate_clicked')) and (not tz_val)
    
    # Check if field was auto-populated (has value and place was checked)
    is_auto_populated = bool(tz_val and st.session_state.get('last_place_checked', ''))
    
    # Always disable UTC field until place is entered (force proper workflow)
    should_disable = not place_val or is_auto_populated
    
    if is_auto_populated:
        render_label('UTC offset (auto-detected) <span style="color:green">✓</span>', False)
    elif not place_val:
        render_label('UTC offset (enter Place of Birth first)', False)
    else:
        # Auto-detection failed, field is editable but still required
        render_label('UTC offset (manual entry required) <span style="color:red">*</span>', tz_err)
    
    tz_override = st.text_input("UTC Offset", key="tz_input", label_visibility="collapsed", disabled=should_disable)

st.write("")
# === End reorganized form layout ===

api_key = st.secrets.get("GEOAPIFY_API_KEY","")

# Center the Generate Kundali button
col1, col2, col3 = st.columns([1, 1, 1])
with col2:
    generate_clicked = st.button("Generate Kundali", key="gen_btn")
    if generate_clicked:
        st.session_state['generate_clicked'] = True
        st.session_state['submitted'] = True
        st.rerun()  # Immediate rerun to show validation

# --- Validation gate computed on rerun after click ---
can_generate = False
if generate_clicked or st.session_state.get('submitted'):
    # Set submitted state for error highlighting
    st.session_state['submitted'] = True
    
    # Use session state values (more reliable after rerun)
    _name = (st.session_state.get('name_input') or '').strip()
    _place = (st.session_state.get('place_input') or '').strip()
    _tz = (st.session_state.get('tz_input') or '').strip()
    _dob = st.session_state.get('dob_input', datetime.date.today())  # Use today as default
    _tob = st.session_state.get('tob_input', datetime.time(12, 0))  # Use 12:00 as default
    
    
    any_err = False
    
    # Check all required fields
    if not _name or not _place or not _tz or _dob is None or _tob is None:
        any_err = True
    else:
        try:
            _tzv = float(_tz)
            if _tzv < -12 or _tzv > 14:
                any_err = True
        except Exception as e:
            any_err = True
    
    if any_err:
        # Error message perfectly centered below the Generate button
        st.markdown(
            """<div style='
                display: flex; 
                justify-content: center; 
                width: 100%; 
                margin-top: 10px;
            '>
                <div style='
                    color: #c1121f; 
                    font-weight: 700; 
                    text-align: center;
                    padding: 8px 0;
                '>
                    Please fix the highlighted fields above.
                </div>
            </div>""", 
            unsafe_allow_html=True
        )
    else:
        can_generate = True
        # Clear previous generation flag to ensure clean state
        st.session_state['generation_completed'] = False

# Show download button only if Kundali was generated in this session

if can_generate:
    # key presence
    api_key = st.secrets.get("GEOAPIFY_API_KEY", "")
    if not api_key:
        st.error("Geoapify key missing. Add GEOAPIFY_API_KEY in Secrets.")
        st.stop()
    
    try:
            # Use the validated variables from session state
            name = _name
            place = _place
            dob = _dob
            tob = _tob
            tz_override = _tz
            
            lat, lon, disp = geocode(place, api_key)
            
            dt_local = datetime.datetime.combine(dob, tob).replace(tzinfo=None)
            used_manual = False
            if tz_override.strip():
                tz_hours = float(tz_override)
                dt_utc = dt_local - datetime.timedelta(hours=tz_hours)
                tzname = f"UTC{tz_hours:+.2f} (manual)"
                used_manual = True
            else:
                tzname, tz_hours, dt_utc = tz_from_latlon(lat, lon, dt_local)

            jd, ay, sidelons = sidereal_positions(dt_utc)
            lagna_sign, asc_sid = ascendant_sign(jd, lat, lon, ay)
            nav_lagna_sign = navamsa_sign_from_lon_sid(asc_sid)

            df_positions = positions_table_no_symbol(sidelons)

            ORDER = ['Ke','Ve','Su','Mo','Ma','Ra','Ju','Sa','Me']
            YEARS = {'Ke':7,'Ve':20,'Su':6,'Mo':10,'Ma':7,'Ra':18,'Ju':16,'Sa':19,'Me':17}

            def moon_balance_days(moon_sid):
                NAK=360.0/27.0; part = moon_sid % 360.0; ni = int(part // NAK); pos = part - ni*NAK
                md_lord = ORDER[ni % 9]; frac = pos/NAK; remaining_days = YEARS[md_lord]*(1 - frac)*YEAR_DAYS
                return md_lord, remaining_days

            def build_mahadashas_days_utc(birth_utc_dt, moon_sid):
                md_lord, rem_days = moon_balance_days(moon_sid); end_limit = birth_utc_dt + datetime.timedelta(days=100*YEAR_DAYS)
                segments=[]; birth_md_start = birth_utc_dt; birth_md_end = min(birth_md_start + datetime.timedelta(days=rem_days), end_limit)
                segments.append({"planet": md_lord, "start": birth_md_start, "end": birth_md_end, "days": rem_days})
                idx = (ORDER.index(md_lord) + 1) % 9; t = birth_md_end
                while t < end_limit:
                    L = ORDER[idx]; dur_days = YEARS[L]*YEAR_DAYS; end = min(t + datetime.timedelta(days=dur_days), end_limit)
                    segments.append({"planet": L, "start": t, "end": end, "days": dur_days}); t = end; idx = (idx + 1) % 9
                return segments
            
            md_segments_utc = build_mahadashas_days_utc(dt_utc, sidelons['Mo'])

            def age_years(birth_dt_local, end_utc):
                local_end = _utc_to_local(end_utc, tzname, tz_hours, used_manual)
                days = (local_end.date() - birth_dt_local.date()).days
                return int(days // YEAR_DAYS)

            df_md = pd.DataFrame([
                {"ग्रह": HN[s["planet"]],
                 "समाप्ति तिथि": _utc_to_local(s["end"], tzname, tz_hours, used_manual).strftime("%d-%m-%Y"),
                 "आयु (वर्ष)": age_years(dt_local, s["end"])}
                for s in md_segments_utc
            ])

            now_utc = datetime.datetime.utcnow()
            rows_an = next_antar_in_days_utc(now_utc, md_segments_utc, days_window=365*10)
            df_an = pd.DataFrame([
                {"महादशा": HN[r["major"]], "अंतरदशा": HN[r["antar"]],
                 "तिथि": _utc_to_local(r["end"], tzname, tz_hours, used_manual).strftime("%d-%m-%Y")}
                for r in rows_an
            ]).head(5)

            img_lagna = render_north_diamond(size_px=800, stroke=3)
            img_nav   = render_north_diamond(size_px=800, stroke=3)
            # ===== ENHANCED DOCUMENT SETUP =====
            doc = make_document()
            sec = doc.sections[0]; sec.page_width = Mm(210); sec.page_height = Mm(297)
            margin = Mm(10); sec.left_margin = sec.right_margin = margin; sec.top_margin = Mm(8); sec.bottom_margin = Mm(8)

            # Enhanced document styling
            style = doc.styles['Normal']; style.font.name = LATIN_FONT; style.font.size = Pt(BASE_FONT_PT)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), HINDI_FONT); style._element.rPr.rFonts.set(qn('w:cs'), HINDI_FONT)
            
            # Set subtle page background
            try:
                set_page_background(doc, 'FEFEFE')  # Very light gray background
            except Exception:
                pass

            
            
            
            
            # ===== EXACT LAYOUT MATCH: Top section with Personal Details (left) + MRIDAASTRO (right) =====
            try:
                # Create top header table (2 columns: Personal Details | MRIDAASTRO)
                header_table = doc.add_table(rows=1, cols=2)
                header_table.autofit = False
                left_width_in = 3.6  # inches; Personal Details column
                header_table.columns[0].width = Inches(left_width_in)
                header_table.columns[1].width = Inches(7.5 - left_width_in)  # keep total ~7.5"
  # Right: MRIDAASTRO (adjusted)
                
                # Remove borders from header table
                hdr_tbl = header_table._tbl
                hdr_tblPr = hdr_tbl.tblPr
                hdr_tblBorders = OxmlElement('w:tblBorders')
                for edge in ('top','left','bottom','right','insideH','insideV'):
                    el = OxmlElement(f'w:{edge}')
                    el.set(qn('w:val'), 'nil')
                    hdr_tblBorders.append(el)
                hdr_tblPr.append(hdr_tblBorders)
                
                # LEFT CELL: Personal Details
                left_cell = header_table.rows[0].cells[0]
                left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Personal Details Title
                p_title = left_cell.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_title.paragraph_format.space_before = Pt(0)
                p_title.paragraph_format.space_after = Pt(0)
                r_title = p_title.add_run("व्यक्तिगत विवरण")
                r_title.font.bold = True
                r_title.font.size = Pt(12)
                
                # Create aligned personal details using proper spacing
                details = [
                    ("नाम:", name),
                    ("जन्म तिथि:", dt_local.strftime('%Y-%m-%d')),
                    ("जन्म समय:", dt_local.strftime('%H:%M:%S')),
                    ("स्थान:", place)
                ]
                
                pd_table = left_cell.add_table(rows=len(details), cols=2)
                try:
                    pd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                set_col_widths(pd_table, [1.3, max(1.0, left_width_in - 1.3 - 0.1)])
                for i, (label, value) in enumerate(details):
                    c0 = pd_table.cell(i, 0)
                    c1 = pd_table.cell(i, 1)
                    # Label
                    p0 = c0.paragraphs[0]
                    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p0.paragraph_format.space_before = Pt(0)
                    p0.paragraph_format.space_after = Pt(0)
                    r0 = p0.add_run(str(label))
                    r0.font.bold = True
                    r0.font.size = Pt(10)
                    # Value
                    p1 = c1.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p1.paragraph_format.space_before = Pt(0)
                    p1.paragraph_format.space_after = Pt(0)
                    r1 = p1.add_run(str(value))
                    r1.font.size = Pt(10)
                
                # Add dark orange rounded border around personal details cell using VML
                try:
                    # Create a VML rounded rectangle overlay for the personal details
                    vml_w_pt = int(left_width_in * 72) - 12
                    vml_h_pt = 88
                    vml_content = f'''
                    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                      <w:pPr>
                        <w:spacing w:before="0" w:after="0"/>
                      </w:pPr>
                      <w:r>
                        <w:pict xmlns:v="urn:schemas-microsoft-com:vml">
                          <v:roundrect style="position:absolute;left:0pt;top:0pt;width:{vml_w_pt}pt;height:{vml_h_pt}pt;z-index:-1" 
                                       arcsize="15%" fillcolor="transparent" strokecolor="#CC6600" strokeweight="3pt">
                          </v:roundrect>
                        </w:pict>
                      </w:r>
                    </w:p>'''
                    vml_element = parse_xml(vml_content)
                    left_cell._element.insert(0, vml_element)
                except Exception:
                    # Fallback to regular thick border if VML fails
                    tc = left_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Remove existing borders first
                    existing_borders = tcPr.find(qn('w:tcBorders'))
                    if existing_borders is not None:
                        tcPr.remove(existing_borders)
                    
                    # Add dark orange borders
                    tcBorders = OxmlElement('w:tcBorders')
                    for edge in ('top', 'left', 'bottom', 'right'):
                        el = OxmlElement(f'w:{edge}')
                        el.set(qn('w:val'), 'single')
                        el.set(qn('w:sz'), '18')  # Thick border
                        el.set(qn('w:color'), 'CC6600')  # Dark orange
                        el.set(qn('w:space'), '0')
                        tcBorders.append(el)
                    tcPr.append(tcBorders)
                
                # RIGHT CELL: MRIDAASTRO + Tagline
                right_cell = header_table.rows[0].cells[1]
                
                # MRIDAASTRO - Enhanced font size (48px equivalent = 36pt)
                p_mrid = right_cell.add_paragraph()
                p_mrid.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_mrid = p_mrid.add_run("MRIDAASTRO")
                r_mrid.font.bold = True
                r_mrid.font.size = Pt(36)  # Enhanced from 16pt to 36pt
                r_mrid.font.name = "Cinzel Decorative"
                # Force font type change using XML
                rPr = r_mrid._element.rPr
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        rFonts.set(qn('w:ascii'), 'Cinzel Decorative')
                        rFonts.set(qn('w:hAnsi'), 'Cinzel Decorative')
                        rFonts.set(qn('w:cs'), 'Cinzel Decorative')
                
                # Tagline
                p_tag = right_cell.add_paragraph()
                p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_tag = p_tag.add_run("In the light of the divine, let your soul journey shine.")
                r_tag.italic = True
                r_tag.font.size = Pt(14)  # Enhanced from 10pt to 14pt
                
                # Add some space after header table
                spacer1 = doc.add_paragraph()
                spacer1.paragraph_format.space_after = Pt(6)
                
                # CENTERED DOCUMENT TITLE
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_title_main = title_para.add_run("PERSONAL HOROSCOPE (JANMA KUNDALI)")
                r_title_main.font.bold = True
                r_title_main.font.size = Pt(20)
                
                # Add space after title
                spacer2 = doc.add_paragraph()
                spacer2.paragraph_format.space_after = Pt(4)
                
            except Exception as e:
                # Fallback to simple header
                pass
# ===== End Header Block (simplified & robust) =====
# ===== End Header Block (safe) =====


            # ===== ENHANCED MAIN LAYOUT TABLE =====
            outer = doc.add_table(rows=1, cols=2); outer.autofit=False
            right_width_in = 3.70; outer.columns[0].width = Inches(3.70); outer.columns[1].width = Inches(3.70)

            CHART_W_PT = int(right_width_in * 72 - 10)
            CHART_H_PT = int(CHART_W_PT * 0.80)
            ROW_HEIGHT_PT = int(CHART_H_PT + 14)
            
            # Remove outer borders, keep only internal divider with dark orange
            tbl = outer._tbl; tblPr = tbl.tblPr; tblBorders = OxmlElement('w:tblBorders')
            # Remove all outer borders
            for edge in ('top','left','bottom','right'):
                el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'nil'); tblBorders.append(el)
            # Keep internal vertical divider with dark orange
            for edge in ('insideV',):
                el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'), 'D2691E'); tblBorders.append(el)
            # Remove horizontal internal borders
            for edge in ('insideH',):
                el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'nil'); tblBorders.append(el)
            tblPr.append(tblBorders)
            
            # Add subtle table shading
            try:
                tblPr = outer._tbl.tblPr
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FDFDFD')  # Very light background
                tblPr.append(shd)
            except Exception:
                pass

            left = outer.rows[0].cells[0]
            # ===== MODERN PERSONAL DETAILS SECTION WITH UNIFIED ROUNDED BOX =====            
            # Get place display value
            try:
                place_disp = disp
            except Exception:
                place_disp = place if 'place' in locals() else ''
            
            # Personal details are now in the header section above, no need for duplicate
            # Original planetary positions section
            # h1 = left.add_paragraph("ग्रह स्थिति"); _apply_hindi_caption_style(h1, size_pt=11, underline=True, bold=True)
            create_cylindrical_section_header(left, "ग्रह स्थिति", width_pt=260)
            
            # === COMPLETELY REWRITTEN FIRST TABLE: ग्रह स्थिति ===
            # Create table with exact 5 columns for clean structure
            t1 = left.add_table(rows=1, cols=5)
            t1.autofit = False  # Disable autofit to prevent conflicts
            
            # Set headers manually to ensure correct order
            headers = ["ग्रह", "राशि", "अंश", "नक्षत्र", "उप‑नक्षत्र"]
            for i, header in enumerate(headers):
                t1.rows[0].cells[i].text = header
            
            # Add data rows with clean structure
            for _, row in df_positions.iterrows():
                new_row = t1.add_row()
                new_row.cells[0].text = str(row["ग्रह"]) if pd.notna(row["ग्रह"]) else ""
                new_row.cells[1].text = str(row["राशि"]) if pd.notna(row["राशि"]) else ""
                new_row.cells[2].text = str(row["अंश"]) if pd.notna(row["अंश"]) else ""
                new_row.cells[3].text = str(row["नक्षत्र"]) if pd.notna(row["नक्षत्र"]) else ""
                new_row.cells[4].text = str(row["उप‑नक्षत्र"]) if pd.notna(row["उप‑नक्षत्र"]) else ""
                
                # Center align all data cells
                for cell in new_row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply styling and formatting
            center_header_row(t1)
            set_table_font(t1, pt=BASE_FONT_PT)
            add_table_borders(t1, size=6)
            apply_premium_table_style(t1)
            
            # Set proper column widths AFTER creating structure
            set_col_widths(t1, [0.70, 0.55, 0.85, 0.80, 0.80])
            
            # Left align ONLY the header cell of the last column (उप‑नक्षत्र)
            for p in t1.rows[0].cells[-1].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT


            # Original Mahadasha section
            # h2 = left.add_paragraph("विंशोत्तरी महादशा"); _apply_hindi_caption_style(h2, size_pt=11, underline=True, bold=True); h2.paragraph_format.keep_with_next = True; h2.paragraph_format.space_after = Pt(2)
            create_cylindrical_section_header(left, "विंशोत्तरी महादशा", width_pt=260)
            t2 = left.add_table(rows=1, cols=len(df_md.columns)); t2.autofit=True
            for i,c in enumerate(df_md.columns): t2.rows[0].cells[i].text=c
            for _,row in df_md.iterrows():
                r=t2.add_row().cells
                for i,c in enumerate(row): 
                    # Clean data handling - avoid NaN and empty values
                    val = str(c) if pd.notna(c) and str(c).strip() else ""
                    r[i].text = val
                    # Ensure proper cell alignment
                    for p in r[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            center_header_row(t2); set_table_font(t2, pt=BASE_FONT_PT); add_table_borders(t2, size=6)
            apply_premium_table_style(t2)  # Apply orange headers and alternating grey rows
            set_col_widths(t2, [1.20, 1.50, 1.00])

            # Original Antardasha section
            # h3 = left.add_paragraph("महादशा / अंतरदशा"); _apply_hindi_caption_style(h3, size_pt=11, underline=True, bold=True)
            create_cylindrical_section_header(left, "महादशा / अंतरदशा", width_pt=260)
            t3 = left.add_table(rows=1, cols=len(df_an.columns)); t3.autofit=True
            for i,c in enumerate(df_an.columns): t3.rows[0].cells[i].text=c
            for _,row in df_an.iterrows():
                r=t3.add_row().cells
                for i,c in enumerate(row): 
                    # Clean data handling - avoid NaN and empty values
                    val = str(c) if pd.notna(c) and str(c).strip() else ""
                    r[i].text = val
                    # Ensure proper cell alignment
                    for p in r[i].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            center_header_row(t3); set_table_font(t3, pt=BASE_FONT_PT); add_table_borders(t3, size=6)
            apply_premium_table_style(t3)  # Apply orange headers and alternating grey rows
            set_col_widths(t3, [1.30, 1.40, 1.00])  # Adjusted column widths for better alignment
            compact_table_paragraphs(t3)  # Move after styling to prevent border conflicts

            # One-page: place Pramukh Bindu under tables (left column) to free right column for charts
            try:
                add_pramukh_bindu_section(left, sidelons, lagna_sign, dt_utc)
                add_phalit_section(left, rows=12)  # Reduced rows to prevent overlapping
            except Exception:
                pass
            right = outer.rows[0].cells[1]

            # Ensure the OUTER right cell has zero inner margins so the kundali touches the cell borders
            try:
                right_tcPr = right._tc.get_or_add_tcPr()
                right_tcMar = right_tcPr.find('./w:tcMar')
                if right_tcMar is None:
                    right_tcMar = OxmlElement('w:tcMar')
                    right_tcPr.append(right_tcMar)
                for side in ('top','left','bottom','right'):
                    el = OxmlElement(f'w:{side}')
                    el.set(qn('w:w'),'0')
                    el.set(qn('w:type'),'dxa')
                    right_tcMar.append(el)
            except Exception:
                pass

            kt = right.add_table(rows=2, cols=1); kt.autofit=False; kt.columns[0].width = Inches(right_width_in)

            # remove cell padding for chart table to let kundali touch the cell borders
            try:
                tcPr = kt._tbl.tblPr
                tblCellMar = OxmlElement('w:tblCellMar')
                for side in ('top','left','bottom','right'):
                    el = OxmlElement(f'w:{side}')
                    el.set(qn('w:w'),'0')
                    el.set(qn('w:type'),'dxa')
                    tblCellMar.append(el)
                tcPr.append(tblCellMar)
            except Exception:
                pass
            # Compact right-cell paragraph spacing
            try:
                for p in right.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            kt.autofit = False
            kt.columns[0].width = Inches(right_width_in)
            for row in kt.rows: row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY; row.height = Pt(ROW_HEIGHT_PT)
            

            # Original Lagna chart title
            cell1 = kt.rows[0].cells[0]; cap1 = cell1.add_paragraph("लग्न कुंडली")
            cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER; _apply_hindi_caption_style(cap1, size_pt=11, underline=True, bold=True); cap1.paragraph_format.space_before = Pt(2); cap1.paragraph_format.space_after = Pt(2)
            p1 = cell1.add_paragraph(); p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(4)
            # Lagna chart with planets in single box per house
            rasi_house_planets = build_rasi_house_planets_marked(sidelons, lagna_sign)
            p1._p.addnext(kundali_with_planets(size_pt=CHART_W_PT, lagna_sign=lagna_sign, house_planets=rasi_house_planets))

            # Original Navamsa chart title - Enhanced styling for visibility
            cell2 = kt.rows[1].cells[0]; cap2 = cell2.add_paragraph("नवांश कुंडली")
            cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Ensure the title is visible with proper formatting
            if cap2.runs:
                run = cap2.runs[0]
            else:
                run = cap2.add_run("नवांश कुंडली")
            run.bold = True
            run.underline = True  
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(139, 69, 19)  # Saddle brown color
            cap2.paragraph_format.space_before = Pt(2); cap2.paragraph_format.space_after = Pt(2)
            p2 = cell2.add_paragraph(); p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(4)
            nav_house_planets = build_navamsa_house_planets_marked(sidelons, nav_lagna_sign)
            p2._p.addnext(kundali_with_planets(size_pt=CHART_W_PT, lagna_sign=nav_lagna_sign, house_planets=nav_house_planets))
            # (प्रमुख बिंदु moved to row 2 of outer table)
            # Ensure content goes below chart shape - single spacing paragraph
            cell2.add_paragraph("").paragraph_format.space_after = Pt(0)
            # (Pramukh Bindu moved above charts)

            out = BytesIO();
            # APPLY_ZERO_MARGINS_BEFORE_SAVE
            try:
                for tbl in doc.tables:
                    zero_table_cell_margins(tbl)
            except Exception:
                pass
            compact_document_spacing(doc)
            doc.save(out); out.seek(0)
            # Store document data in session state for download button
            st.session_state['kundali_doc'] = out.getvalue()
            st.session_state['kundali_filename'] = f"{sanitize_filename(name)}_Horoscope.docx"
            st.session_state['generation_completed'] = True

    except Exception as e:
        st.error(f"Error generating Kundali: {str(e)}")
        import traceback
        st.code(traceback.format_exc())

# Show download button centered below Generate button after validation
if (st.session_state.get('kundali_doc') and 
    st.session_state.get('generation_completed') and
    st.session_state.get('submitted') and  # User must have clicked Generate
    can_generate):  # AND current form is still valid
    
    # Center the download button like the Generate button
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        st.download_button(
            "📥 Download Kundali (DOCX)", 
            st.session_state['kundali_doc'], 
            file_name=st.session_state.get('kundali_filename', 'Horoscope.docx'),
            type="primary",
            key="download_button_main"
        )


if __name__=='__main__':
    main()
